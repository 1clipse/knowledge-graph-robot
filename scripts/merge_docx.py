#!/usr/bin/env python3
"""Merge two docx files into one Chinese document"""
from __future__ import annotations

import copy
from datetime import datetime
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

REPORT = Path.home() / "Desktop" / "QLoRA_Fine-tuning_Report_2026-05-22.docx"
MANUAL = Path.home() / "Desktop" / "QLoRA微调操作手册.docx"
OUTPUT = Path.home() / "Desktop" / "Qwen2.5-7B_QLoRA微调完整指南_2026-05-22.docx"

report = Document(str(REPORT))
manual = Document(str(MANUAL))

merged = Document()

# ── Style ──
style = merged.styles["Normal"]
font = style.font
font.name = "Microsoft YaHei"
font.size = Pt(11)

# ── Cover title ──
title = merged.add_heading("Qwen2.5-7B QLoRA 微调完整指南", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
merged.add_paragraph(f"整合日期: {datetime.now().strftime('%Y-%m-%d')}")
merged.add_paragraph("内容来源: QLoRA微调操作手册 + 2026-05-22 训练报告")
merged.add_page_break()

# ── Part 1: Operation Manual (Chinese) ──
merged.add_heading("第一部分：QLoRA 微调操作手册", level=1)
merged.add_paragraph("以下内容来自「QLoRA微调操作手册.docx」，涵盖微调基础知识和操作流程。")

for element in manual.element.body:
    merged.element.body.append(copy.deepcopy(element))

merged.add_page_break()

# ── Part 2: Training Report (Chinese translation) ──
merged.add_heading("第二部分：2026-05-22 微调训练实战报告", level=1)

h = merged.add_heading("1. 概述", level=1)
merged.add_paragraph(
    "使用 QLoRA 技术对 Qwen2.5-7B-Instruct 模型进行微调，"
    "训练数据来自 6 张工业机器人 DWG 图纸，"
    "目标是让模型能够从工业设计文档中抽取知识图谱所需的实体和关系。"
)

merged.add_heading("1.1 硬件环境", level=2)
merged.add_paragraph("GPU: 2x Tesla P100-SXM2-16GB（仅使用 GPU 0 进行训练）")
merged.add_paragraph("服务器: x86_64, ZeroTier 内网 IP 10.117.29.24")
merged.add_paragraph("训练耗时: 约 6.5 分钟 / 3 轮")

merged.add_heading("1.2 训练数据", level=2)
merged.add_paragraph("数据来源: C:\\Users\\Knightz\\Desktop\\train_dwg 下的 6 张 DWG 图纸")
merged.add_paragraph("最终数据集: 27 条（3 条手工标注 + 24 条数据增强）")
merged.add_paragraph("实体类型: 13 种（Robot, Manufacturer, Component, Reducer, "
                     "ServoMotor, Controller, Sensor, ApplicationScenario, "
                     "Process, EndEffector, Standard, Material, Software）")
merged.add_paragraph("关系类型: 10 种")

# Pipeline
merged.add_heading("2. 完整管线", level=1)

merged.add_heading("步骤1：增强标注 (enhanced_label.py)", level=2)
merged.add_paragraph("润色 DWG 描述 — 去除 %%c、%%p 等 CAD 乱码，改写为自然语言")
merged.add_paragraph("两阶段抽取 — 先抽取实体，再基于实体抽取关系（比单阶段更准确）")
merged.add_paragraph("Schema 校验 — 对照工业机器人本体 Schema 验证实体和关系合法性")
merged.add_paragraph("质量评分 — good / ok / review 三档")
merged.add_paragraph("数据增强 — 改写描述生成变体，扩充训练集")
merged.add_paragraph("输出: enhanced_labeled.json + LLaMA-Factory 格式训练数据")

merged.add_heading("步骤2：自动清洗 (clean_labeled.py)", level=2)
merged.add_paragraph("去除 CAD 图元 — 圆、圆弧、线段、图块、中心标记等")
merged.add_paragraph("去除噪音实体 — 标准材料、尺寸信息、标注文字、注释")
merged.add_paragraph("修正关系方向 — contains（容器→被包含）、performs_process（零件→工艺）")
merged.add_paragraph("去重 — 双向重复关系合并")
merged.add_paragraph("实体名规范化 — 4-M → M4螺纹孔")
merged.add_paragraph("清洗结果: 21 处修正")

merged.add_heading("步骤3：QLoRA 训练 (03_qlora_train.py)", level=2)
merged.add_paragraph("基座模型: Qwen2.5-7B-Instruct（保存在 /media/z/data/models/，未修改）")
merged.add_paragraph("量化方式: 4-bit nf4 (bitsandbytes)")
merged.add_paragraph("LoRA 配置: rank=16, alpha=32, target_modules=all-linear")
merged.add_paragraph("训练参数: batch=1×2, max_seq=1024, lr=2e-4, epochs=3")
merged.add_paragraph("可训练参数: 40M / 4.4B (0.92%)")
merged.add_paragraph("Loss: 2.4 → 0.1（3 轮）")

merged.add_heading("步骤4：合并与部署 (04_merge_and_deploy.sh)", level=2)
merged.add_paragraph("将 LoRA 适配器合并为完整模型，保存到独立路径（不覆盖基座模型）")
merged.add_paragraph("合并后模型: /data/finetune/output/qwen2.5-7b-kg-robot-merged/")
merged.add_paragraph("更新 P100 API 配置: MODEL_PATH + MAX_MEMORY + device_map")

# Key fixes
merged.add_heading("3. 关键问题与修复", level=1)

fixes = [
    ("测试 Prompt 不匹配",
     "训练时使用约 800 字符的指令（含实体/关系类型列表 + JSON 格式规则 + 5条抽取规则），"
     "测试时只用了一行通用系统消息（约 50 字符），导致模型输出自然语言而非 JSON。"
     "修复: test_inference() 中使用与训练一致的完整 instruction。"),
    ("P100 GPU 配置错误",
     "MAX_MEMORY={0:'12GiB',1:'12GiB'} 配合 device_map='auto' 导致 "
     "CUDA error: invalid device ordinal。"
     "修复: 改为 MAX_MEMORY={0:'15GiB'} + device_map='cuda:0'，单卡运行。"),
    ("数据增强死循环",
     "向 samples 列表追加增强样本时同时在遍历该列表，导致无限循环。"
     "修复: 先对原始样本做快照遍历，新样本收集到独立列表，循环结束后再合并。"),
    ("DXF 规则引擎 vs LLM 抽取路径",
     "通过 /ingest/file 上传 DWG 走的是 DXF 规则引擎，抽取出的是 CAD 图元 "
     "（SW_NOTE_0、SW_CENTERMARKSYMBOL_0），不是语义知识。"
     "要使用微调模型，必须通过 /ingest/text 上传 .txt 描述文件。"
     "工具脚本: batch_upload_text.py + 08_upload_text.bat"),
]

for t, desc in fixes:
    merged.add_heading(t, level=2)
    merged.add_paragraph(desc)

# Integration
merged.add_heading("4. 知识图谱集成", level=1)
merged.add_paragraph(
    "调用链路: 用户 .txt 描述 → KG /api/v1/ingest/text → "
    "LLMExtractor → P100 微调模型 (http://10.117.29.24:5200/v1) → Neo4j"
)
merged.add_paragraph(
    "问答页面 http://localhost:8000/chat 从 Neo4j 检索微调模型抽取的实体和关系，"
    "由 LLM 生成自然语言答案。"
)

# File locations
merged.add_heading("5. 关键文件位置", level=1)
files = [
    ("Windows 本地", "E:\\Knowledge Graph_robot\\scripts\\finetune\\"),
    ("P100 服务器", "/data/finetune/"),
    ("基座模型（未修改）", "/media/z/data/models/Qwen2.5-7B-Instruct"),
    ("合并后模型", "/data/finetune/output/qwen2.5-7b-kg-robot-merged/"),
    ("训练数据", "data/handcrafted_examples.json (27条)"),
    ("增强标注脚本", "enhanced_label.py"),
    ("自动清洗脚本", "clean_labeled.py"),
    ("QLoRA 训练脚本", "03_qlora_train.py"),
    ("合并部署脚本", "04_merge_and_deploy.sh"),
    ("文本入库脚本", "batch_upload_text.py"),
    ("桌面批处理文件", "04_enhanced_label.bat ~ 08_upload_text.bat"),
]
for name, path in files:
    p = merged.add_paragraph()
    p.add_run(f"{name}: ").bold = True
    p.add_run(path)

# P100 API config notes
merged.add_heading("6. P100 API 配置注意事项", level=1)
merged.add_paragraph("MAX_MEMORY = {0: '15GiB'} — 只用 GPU 0")
merged.add_paragraph("device_map = 'cuda:0' — 显式指定单卡，不使用 'auto'")
merged.add_paragraph("CUDA_VISIBLE_DEVICES=0 — 环境变量限制单卡")
merged.add_paragraph("测试 prompt 必须与训练 instruction 一致")

# Training results
merged.add_heading("7. 训练结果汇总", level=1)
merged.add_paragraph("训练数据: 27 条（3 handcrafted + 24 augmented）")
merged.add_paragraph("Loss: 2.4 → 0.1")
merged.add_paragraph("训练时间: 6.5 分钟")
merged.add_paragraph("测试输出: JSON 格式正确，实体类型正确")
merged.add_paragraph("可训练参数: 40M / 4.4B (0.92%)")
merged.add_paragraph("知识图谱入库: 6/6 DWG 全部成功")
merged.add_paragraph("问答验证: 材质查询、工艺查询均正确返回")

# Save
merged.save(str(OUTPUT))
print(f"Saved: {OUTPUT}")
