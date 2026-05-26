#!/usr/bin/env python3
"""Add data augmentation pipeline details to v3 docx"""
from docx import Document
from docx.shared import Pt
from lxml import etree

INPUT = r"C:\Users\Knightz\Desktop\Qwen2.5-7B_QLoRA微调完整指南_v3.docx"
OUTPUT = r"C:\Users\Knightz\Desktop\v3_temp.docx"
doc = Document(INPUT)

# Find the paragraph after "步骤2：自动清洗" to insert augmentation section
target = None
target_idx = None
for i, p in enumerate(doc.paragraphs):
    if "enhanced_labeled.json + LLaMA-Factory" in p.text:
        target = p
        target_idx = i
        break

if target is None:
    print("ERROR: target paragraph not found")
    exit(1)

# Find parent element and position
parent = target._element.getparent()
elem_idx = list(parent).index(target._element)

def make_heading(text, level):
    h = doc.add_heading(text, level=level)
    return h._element

def make_para(text, bold_prefix=None):
    p = doc.add_paragraph()
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p._element

elements = []

# Insert after target
elements.append(make_para(""))

h = doc.add_heading("数据增强详情", level=2)
elements.append(h._element)

elements.append(make_para(
    "enhanced_label.py 的 --augment 参数通过改写描述文本生成训练变体，"
    "每条原始样本生成多个语义等价但表述不同的副本，提升模型泛化能力。"
))

elements.append(make_para("数据管线各阶段数量："))
pipeline = [
    ("handcrafted_examples.json", "3 条", "原始手工标注"),
    ("enhanced_labeled.json", "6 条", "LLM 增强标注（润色 + Schema 校验）"),
    ("enhanced_labeled_cleaned.json", "6 条", "自动清洗（去 CAD 图元/噪音/去重）"),
    ("enhanced_labeled_cleaned_augmented.json", "24 条", "数据增强（--augment 3, 6x4=24）"),
    ("enhanced_labeled_cleaned_augmented_llamafactory.json", "24 条", "LLaMA-Factory 训练格式"),
]
for fname, count, desc in pipeline:
    elements.append(make_para(f"  {fname}: {count} — {desc}"))

elements.append(make_para(""))
elements.append(make_para(
    "增强策略: 对每条清洗后的样本改写 3 个变体（--augment 3），"
    "生成时保持实体/关系标注不变，仅改写原文描述。6 条 -> 24 条，4 倍扩充。"
))
elements.append(make_para(
    "命令: python enhanced_label.py --dir <目录> --augment 3"
))

# Insert in reverse order after target
for elem in reversed(elements):
    parent.insert(elem_idx + 1, elem)

doc.save(OUTPUT)
print(f"Updated: {OUTPUT}")
