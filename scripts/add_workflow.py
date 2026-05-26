#!/usr/bin/env python3
"""Add workflow flowchart to the merged docx"""
from __future__ import annotations

from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

DOC = Path.home() / "Desktop" / "Qwen2.5-7B_QLoRA微调完整指南_2026-05-22.docx"
OUTPUT = Path.home() / "Desktop" / "Qwen2.5-7B_QLoRA微调完整指南_v2.docx"
doc = Document(str(DOC))

# ── Helper: create colored table cell ──
def set_cell_shading(cell, color):
    """Set cell background color"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>')
    tcPr.append(shading)

def set_cell_text(cell, text, bold=False, size=9, color=None, align=None):
    """Set cell text with formatting"""
    cell.text = ""
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)

def add_arrow_row(table, cols, text=""):
    """Add an arrow/down indicator row"""
    row = table.add_row()
    for i in range(cols):
        set_cell_shading(row.cells[i], "FFFFFF")
        set_cell_text(row.cells[i], "", size=6)
    # Merge all cells for the arrow
    if cols > 1:
        row.cells[0].merge(row.cells[-1])
    set_cell_text(row.cells[0], text if text else "    ▼", size=10, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

# ── Insert workflow section before the end ──
# Find the last element to append after
# Actually just append at the end since we already have all content

doc.add_page_break()
doc.add_heading("第三部分：完整工作流（端到端）", level=1)
doc.add_paragraph(
    "以下流程图覆盖从原始 DWG 文件到知识图谱问答的完整链路，"
    "包括数据准备、模型训练、部署、入库和查询五个阶段。"
)

# ═══ STAGE 0: Overview ═══
doc.add_heading("流程总览", level=2)

overview = doc.add_table(rows=1, cols=5, style="Table Grid")
overview.alignment = WD_TABLE_ALIGNMENT.CENTER

stage_colors = ["4472C4", "ED7D31", "A5A5A5", "70AD47", "FFC000"]
stage_names = [
    "阶段一\n数据准备\n(Windows)",
    "阶段二\n模型训练\n(P100服务器)",
    "阶段三\n模型部署\n(P100服务器)",
    "阶段四\n知识入库\n(Windows→KG)",
    "阶段五\n问答查询\n(浏览器)"
]

for i, (name, color) in enumerate(zip(stage_names, stage_colors)):
    set_cell_shading(overview.rows[0].cells[i], color)
    set_cell_text(overview.rows[0].cells[i], name, bold=True, size=10,
                  color=(255, 255, 255), align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_paragraph()

# ═══ STAGE 1: Data Preparation ═══
doc.add_heading("阶段一：数据准备（Windows 本地）", level=2)

t1 = doc.add_table(rows=1, cols=1, style="Table Grid")
t1.alignment = WD_TABLE_ALIGNMENT.CENTER
set_cell_shading(t1.rows[0].cells[0], "E8F0FE")

stage1_text = """❶ 原始 DWG 文件 (6张)
  C:\\Users\\Knightz\\Desktop\\train_dwg\\
    CHK-3前爪法兰.DWG    CHK-3前爪法兰侧盖.DWG
    CHK-3前爪法兰侧钣金盖.DWG    CHX-3大手臂摆线减速机安装法兰.DWG
    CHX-3底盘旋转涡轮箱.DWG    02 机器人的自冲铆接结构性设计 2d 发.DWG

    ▼

❷ 提取 .txt 描述文本
  使用大模型对每张 DWG 生成自然语言描述
  描述内容: 零件名称、材质、工艺、装配关系、技术标准
  输出: 6 个 .txt 文件（同名）

    ▼

❸ 增强标注 (enhanced_label.py)
  命令: python enhanced_label.py --dir <目录> --augment 3
  步骤:
    a. 润色描述 → 去除 %%c/%%p CAD 乱码，自然语言改写
    b. 两阶段抽取 → 先抽实体（13类），再基于实体抽关系（10类）
    c. Schema 校验 → 验证实体/关系是否符合工业机器人本体
    d. 质量评分 → good / ok / review
  输出: enhanced_labeled.json

    ▼

❹ 自动清洗 (clean_labeled.py)
  命令: python clean_labeled.py
  处理:
    a. 去除 CAD 图元实体（圆/圆弧/线段/图块/中心标记）
    b. 去除噪音实体（标准材料/尺寸信息/标注文字）
    c. 修正关系方向（contains/performs_process）
    d. 去重双向关系
    e. 规范化实体名（4-M → M4螺纹孔）
  结果: 21 处修正

    ▼

❺ 数据增强
  对每条已清洗的标注数据生成 N 个描述变体
  保留相同的实体/关系输出
  输出: handcrafted_examples.json (27条: 3原始 + 24增强)

    ▼

❻ 上传到 P100 服务器
  命令: 06_upload_and_train.bat
  scp data/ 到 P100:/data/finetune/data/"""
set_cell_text(t1.rows[0].cells[0], stage1_text, size=9)

doc.add_paragraph()

# ═══ STAGE 2: Training ═══
doc.add_heading("阶段二：模型训练（P100 服务器）", level=2)

t2 = doc.add_table(rows=1, cols=1, style="Table Grid")
t2.alignment = WD_TABLE_ALIGNMENT.CENTER
set_cell_shading(t2.rows[0].cells[0], "FDE8D0")

stage2_text = """❶ 连接 P100 服务器
  SSH: ssh z@10.117.29.24 (密码: z)
  服务器: x86_64, 2x Tesla P100-SXM2-16GB

    ▼

❷ 启动 QLoRA 训练 (03_qlora_train.py)
  命令: CUDA_VISIBLE_DEVICES=0 python 03_qlora_train.py train

  训练配置:
    基座模型: Qwen2.5-7B-Instruct (HuggingFace)
    量化: 4-bit nf4 (bitsandbytes)
    LoRA: rank=16, alpha=32, target_modules=all-linear
    优化器: paged_adamw_8bit
    超参: batch=1×2, max_seq=1024, lr=2e-4, epochs=3

  训练结果:
    数据量: 27 条
    Loss: 2.4 → 0.1
    耗时: ~6.5 分钟
    可训练参数: 40M / 4.4B (0.92%)

    ▼

❸ 测试推理
  命令: python 03_qlora_train.py test
  加载 LoRA adapter → 运行测试用例
  验证: JSON 格式正确、实体类型正确、关系方向正确"""
set_cell_text(t2.rows[0].cells[0], stage2_text, size=9)

doc.add_paragraph()

# ═══ STAGE 3: Deploy ═══
doc.add_heading("阶段三：模型部署（P100 服务器）", level=2)

t3 = doc.add_table(rows=1, cols=1, style="Table Grid")
t3.alignment = WD_TABLE_ALIGNMENT.CENTER
set_cell_shading(t3.rows[0].cells[0], "E8E8E8")

stage3_text = """❶ 合并 LoRA → 完整模型 (04_merge_and_deploy.sh)
  将 LoRA adapter 权重合并到基座模型
  输出路径: /data/finetune/output/qwen2.5-7b-kg-robot-merged/
  基座模型保持不变: /media/z/data/models/Qwen2.5-7B-Instruct

    ▼

❷ 配置 API 服务 (qwen_api.py)
  修改项:
    MODEL_PATH = "/data/finetune/output/qwen2.5-7b-kg-robot-merged/"
    MAX_MEMORY = {0: "15GiB"}          # 只用 GPU 0
    device_map = "cuda:0"              # 显式指定单卡，不用 "auto"

    ▼

❸ 启动 API 服务
  命令: CUDA_VISIBLE_DEVICES=0 python /data/qwen_api.py
  端点: http://10.117.29.24:5200/v1
  验证: curl http://10.117.29.24:5200/health"""
set_cell_text(t3.rows[0].cells[0], stage3_text, size=9)

doc.add_paragraph()

# ═══ STAGE 4: Ingest ═══
doc.add_heading("阶段四：知识入库（Windows → KG → P100）", level=2)

t4 = doc.add_table(rows=1, cols=1, style="Table Grid")
t4.alignment = WD_TABLE_ALIGNMENT.CENTER
set_cell_shading(t4.rows[0].cells[0], "E2EFDA")

stage4_text = """❶ 启动知识图谱后端
  命令: cd E:\\Knowledge Graph_robot && E:\\kg_robot_env\\Scripts\\python.exe -m uvicorn api.app:app --host 0.0.0.0 --port 8000
  端点: http://localhost:8000
  依赖: Neo4j Aura 云数据库 (已配置)

    ▼

❷ 文本入库 (batch_upload_text.py + 08_upload_text.bat)
  遍历 train_dwg/ 下所有 .txt 文件
  对每个文件:
    POST /api/v1/ingest/text {
      "text": "<txt内容>",
      "use_llm": true,           # 使用微调模型
      "use_rule_fallback": false # 不回退到规则引擎
    }

    ▼

❸ LLM 抽取 (KG 后端内部)
  KG 后端接收 text → LLMExtractor 模块
    → 调用 http://10.117.29.24:5200/v1/chat/completions
    → 微调 Qwen2.5-7B 模型
    → 返回结构化 JSON (entities + relations)

    ▼

❹ 写入 Neo4j
  解析 JSON → MERGE 实体节点 + 关系边
  实体标签: Component / Material / Process / Manufacturer / Robot / Standard...
  关系类型: process_material / contains / performs_process / complies_with...
  入库: 6 个 DWG 对应 .txt 全部成功"""
set_cell_text(t4.rows[0].cells[0], stage4_text, size=9)

doc.add_paragraph()

# ═══ STAGE 5: Query ═══
doc.add_heading("阶段五：问答查询（浏览器）", level=2)

t5 = doc.add_table(rows=1, cols=1, style="Table Grid")
t5.alignment = WD_TABLE_ALIGNMENT.CENTER
set_cell_shading(t5.rows[0].cells[0], "FFF2CC")

stage5_text = """❶ 打开问答页面
  浏览器访问: http://localhost:8000/chat

    ▼

❷ 用户提问
  示例问题:
    "减速机安装法兰用什么材质？"
    "前爪法兰侧钣金盖需要什么处理？"
    "这个项目里有多少个CHK-3系列零件？"

    ▼

❸ KG 后端检索
  /api/v1/ask 接口:
    a. 解析用户问题 → 提取关键词
    b. 查询 Neo4j → 检索相关实体和关系
    c. 将检索结果作为上下文发给 LLM

    ▼

❹ LLM 生成答案
  结合 Neo4j 检索结果 + 用户问题 → 自然语言回答
  前端展示: 答案 + 检索到的相关知识卡片

实际效果验证:
  Q: 减速机安装法兰用什么材质？
  A: 减速机安装法兰 CHX-3 大手臂摆线减速机安装法兰.DWG
     使用的材质是 45#钢。  ✓

  Q: 前爪法兰侧钣金盖需要什么处理？
  A: 前爪法兰侧钣金盖需要进行去毛刺处理。  ✓"""
set_cell_text(t5.rows[0].cells[0], stage5_text, size=9)

doc.add_paragraph()

# ═══ Complete Pipeline Summary ═══
doc.add_heading("完整链路图（简化版）", level=2)

pipe = doc.add_table(rows=1, cols=7, style="Table Grid")
pipe.alignment = WD_TABLE_ALIGNMENT.CENTER

pipe_data = [
    ("DWG\n文件", "4472C4"),
    ("txt\n描述", "4472C4"),
    ("增强标注\n+清洗\n+增强", "4472C4"),
    ("QLoRA\n训练", "ED7D31"),
    ("合并部署\nP100 API", "A5A5A5"),
    ("LLM抽取\n→Neo4j", "70AD47"),
    ("问答\n/chat", "FFC000"),
]

for i, (text, color) in enumerate(pipe_data):
    set_cell_shading(pipe.rows[0].cells[i], color)
    set_cell_text(pipe.rows[0].cells[i], text, bold=True, size=9,
                  color=(255, 255, 255), align=WD_ALIGN_PARAGRAPH.CENTER)

# Arrows between columns (using a row below)
arrow_row = pipe.add_row()
for i in range(7):
    set_cell_shading(arrow_row.cells[i], "FFFFFF")
set_cell_text(arrow_row.cells[0], "▶", size=14, bold=True, color=(68, 114, 196), align=WD_ALIGN_PARAGRAPH.CENTER)
set_cell_text(arrow_row.cells[1], "▶", size=14, bold=True, color=(68, 114, 196), align=WD_ALIGN_PARAGRAPH.CENTER)
set_cell_text(arrow_row.cells[2], "▶", size=14, bold=True, color=(237, 125, 49), align=WD_ALIGN_PARAGRAPH.CENTER)
set_cell_text(arrow_row.cells[3], "▶", size=14, bold=True, color=(165, 165, 165), align=WD_ALIGN_PARAGRAPH.CENTER)
set_cell_text(arrow_row.cells[4], "▶", size=14, bold=True, color=(112, 173, 71), align=WD_ALIGN_PARAGRAPH.CENTER)
set_cell_text(arrow_row.cells[5], "▶", size=14, bold=True, color=(255, 192, 0), align=WD_ALIGN_PARAGRAPH.CENTER)
set_cell_text(arrow_row.cells[6], "", size=6)

doc.add_paragraph()

# Key
doc.add_heading("图例", level=2)
legend = doc.add_table(rows=1, cols=5, style="Table Grid")
legend.alignment = WD_TABLE_ALIGNMENT.CENTER

colors = ["4472C4", "ED7D31", "A5A5A5", "70AD47", "FFC000"]
labels = ["数据准备 (Windows)", "模型训练 (P100)", "模型部署 (P100)", "知识入库 (KG)", "问答查询 (浏览器)"]

for i, (c, l) in enumerate(zip(colors, labels)):
    set_cell_shading(legend.rows[0].cells[i], c)
    set_cell_text(legend.rows[0].cells[i], l, bold=True, size=9,
                  color=(255, 255, 255), align=WD_ALIGN_PARAGRAPH.CENTER)

# ── Save ──
doc.save(str(OUTPUT))
print(f"Saved: {OUTPUT}")
