#!/usr/bin/env python3
"""Rewrite docx with LLaMA-Factory complete workflow"""
from __future__ import annotations

import copy
from datetime import datetime
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

MANUAL = Path.home() / "Desktop" / "QLoRA微调操作手册.docx"
OUTPUT = Path.home() / "Desktop" / "Qwen2.5-7B_QLoRA微调完整指南_v3.docx"

manual = Document(str(MANUAL))
doc = Document()

# ── Style ──
style = doc.styles["Normal"]
font = style.font
font.name = "Microsoft YaHei"
font.size = Pt(11)


def set_cell_shading(cell, color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>')
    tcPr.append(shading)


def set_cell_text(cell, text, bold=False, size=9, color=None, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)


# ═══════════════════════════════════════
# COVER
# ═══════════════════════════════════════
title = doc.add_heading("Qwen2.5-7B QLoRA 微调完整指南", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph(f"更新日期: {datetime.now().strftime('%Y-%m-%d')}")
doc.add_paragraph("基于 LLaMA-Factory v0.9.3 框架 | 工业机器人知识图谱抽取")
doc.add_page_break()

# ═══════════════════════════════════════
# PART 1: Original Manual
# ═══════════════════════════════════════
doc.add_heading("第一部分：QLoRA 微调基础知识", level=1)
for element in manual.element.body:
    doc.element.body.append(copy.deepcopy(element))

doc.add_page_break()

# ═══════════════════════════════════════
# PART 2: LLaMA-Factory Complete Pipeline
# ═══════════════════════════════════════
doc.add_heading("第二部分：LLaMA-Factory 完整工作流", level=1)
doc.add_paragraph(
    "本部分覆盖从零开始在 P100 服务器上使用 LLaMA-Factory 框架进行 QLoRA 微调、"
    "模型合并和部署到知识图谱的完整流程。"
)

# ── 2.1 Environment ──
doc.add_heading("2.1 环境准备", level=2)

doc.add_heading("P100 服务器配置", level=3)
env_table = doc.add_table(rows=5, cols=2, style="Table Grid")
env_data = [
    ("GPU", "2x Tesla P100-SXM2-16GB"),
    ("ZeroTier IP", "10.117.29.24"),
    ("Conda 环境", "qwen-api (/data/envs/qwen-api/)"),
    ("Python", "3.10"),
    ("训练方式", "QLoRA: 4-bit nf4 + LoRA (rank=16, alpha=32)"),
]
for i, (k, v) in enumerate(env_data):
    set_cell_shading(env_table.rows[i].cells[0], "2F5496")
    set_cell_shading(env_table.rows[i].cells[1], "D6E4F0")
    set_cell_text(env_table.rows[i].cells[0], k, bold=True, size=10, color=(255, 255, 255))
    set_cell_text(env_table.rows[i].cells[1], v, size=10)

doc.add_paragraph()

doc.add_heading("安装 LLaMA-Factory", level=3)
doc.add_paragraph("在 P100 服务器的 qwen-api 环境中执行：")
install = doc.add_paragraph()
install.add_run("pip install llamafactory -i https://pypi.tuna.tsinghua.edu.cn/simple").font.size = Pt(9)
doc.add_paragraph("安装后验证:")
verify = doc.add_paragraph()
verify.add_run("which llamafactory-cli  # 应输出: /data/envs/qwen-api/bin/llamafactory-cli").font.size = Pt(9)
verify2 = doc.add_paragraph()
verify2.add_run("llamafactory-cli version  # 应显示: LLaMA Factory, version 0.9.3").font.size = Pt(9)

# ── 2.2 Data Preparation ──
doc.add_heading("2.2 数据准备 (Windows 本地)", level=2)

doc.add_heading("步骤 1：DWG → 文本描述", level=3)
doc.add_paragraph(
    "使用大模型对 6 张工业机器人 DWG 图纸生成自然语言描述。"
    "描述内容包括零件名称、材质、工艺、装配关系和技术标准。"
    "输出 6 个 .txt 文件，放在 train_dwg/ 目录下。"
)

doc.add_heading("步骤 2：增强标注", level=3)
doc.add_paragraph("命令:")
cmd = doc.add_paragraph()
cmd.add_run("python enhanced_label.py --dir C:\\Users\\Knightz\\Desktop\\train_dwg --augment 3").font.size = Pt(9)
doc.add_paragraph("流程:")
doc.add_paragraph("a. 润色描述 — 去除 %%c、%%p 等 CAD 乱码，改为自然语言")
doc.add_paragraph("b. 两阶段抽取 — 先抽取实体（13 类），再基于实体抽取关系（10 类）")
doc.add_paragraph("c. Schema 校验 — 对照工业机器人本体验证实体/关系合法性")
doc.add_paragraph("d. 质量评分 — good / ok / review")
doc.add_paragraph("e. 数据增强 — 每条数据生成 3 个描述变体，扩充训练集")
doc.add_paragraph("输出: handcrafted_examples.json + LLaMA-Factory 格式数据")

doc.add_heading("步骤 3：自动清洗", level=3)
doc.add_paragraph("命令:")
cmd = doc.add_paragraph()
cmd.add_run("python clean_labeled.py").font.size = Pt(9)
doc.add_paragraph("处理内容:")
doc.add_paragraph("a. 去除 CAD 图元实体（圆/圆弧/线段/图块/中心标记）")
doc.add_paragraph("b. 去除噪音实体（标准材料/尺寸信息/标注文字）")
doc.add_paragraph("c. 修正关系方向（contains/performs_process）")
doc.add_paragraph("d. 去重双向关系")
doc.add_paragraph("e. 规范化实体名（4-M → M4螺纹孔）")
doc.add_paragraph("最终数据集: 24 条清洗+增强后的 LLaMA-Factory 格式数据")

doc.add_heading("步骤 4：上传到 P100", level=3)
doc.add_paragraph("将 LLaMA-Factory 格式的训练数据上传到 P100:")
cmd = doc.add_paragraph()
cmd.add_run(
    "scp data/enhanced_labeled_cleaned_augmented_llamafactory.json "
    "z@10.117.29.24:/data/finetune/data/kg_robot_llamafactory.json"
).font.size = Pt(9)

# ── 2.3 LLaMA-Factory Configuration ──
doc.add_heading("2.3 LLaMA-Factory 配置", level=2)

doc.add_heading("数据集注册 (dataset_info.json)", level=3)
doc.add_paragraph(
    "LLaMA-Factory 需要通过 dataset_info.json 注册自定义数据集。"
    "该文件放在数据同一目录下（/data/finetune/data/）："
)
config_json = doc.add_paragraph()
config_json.add_run(
    '{\n'
    '  "kg_robot": {\n'
    '    "file_name": "kg_robot_llamafactory.json",\n'
    '    "columns": {\n'
    '      "prompt": "prompt",\n'
    '      "query": "query",\n'
    '      "response": "response",\n'
    '      "system": "system",\n'
    '      "history": "history"\n'
    '    }\n'
    '  }\n'
    '}'
).font.size = Pt(8)

doc.add_heading("训练配置 (train_config.yaml)", level=3)
doc.add_paragraph("在 /data/finetune/ 下创建训练配置文件：")
config_yaml = doc.add_paragraph()
config_yaml.add_run(
    "model_name_or_path: /media/z/data/models/Qwen2.5-7B-Instruct\n"
    "dataset_dir: /data/finetune/data\n"
    "dataset: kg_robot\n"
    "template: qwen\n"
    "finetuning_type: lora\n"
    "stage: sft\n\n"
    "# LoRA\n"
    "lora_rank: 16\n"
    "lora_alpha: 32\n"
    "lora_dropout: 0.05\n"
    "lora_target: all\n\n"
    "# Quantization\n"
    "quantization_method: bitsandbytes\n"
    "quantization_bit: 4\n\n"
    "# Training\n"
    "output_dir: /data/finetune/llamafactory_output\n"
    "overwrite_output_dir: true\n"
    "per_device_train_batch_size: 1\n"
    "gradient_accumulation_steps: 2\n"
    "learning_rate: 2.0e-4\n"
    "lr_scheduler_type: cosine\n"
    "warmup_steps: 5\n"
    "num_train_epochs: 3\n"
    "cutoff_len: 1024\n"
    "max_grad_norm: 0.3\n"
    "logging_steps: 1\n"
    "save_steps: 50\n"
    "save_total_limit: 3\n\n"
    "# Precision (P100: no bf16)\n"
    "fp16: true\n"
    "bf16: false\n\n"
    "# Optimizer\n"
    "optim: paged_adamw_8bit"
).font.size = Pt(8)

doc.add_paragraph()
doc.add_paragraph("参数说明:")
params_t = doc.add_table(rows=8, cols=2, style="Table Grid")
params_data = [
    ("lora_rank=16", "LoRA 秩，越大能力越强但越慢"),
    ("lora_alpha=32", "2x rank，标准设置"),
    ("quantization_bit=4", "4-bit nf4 量化，大幅减少显存"),
    ("learning_rate=2e-4", "学习率，小数据集适中值"),
    ("cutoff_len=1024", "最大序列长度，抽取任务 1024 足够"),
    ("fp16=true / bf16=false", "P100 不支持 bf16，使用 fp16"),
    ("batch=1 x grad_acc=2", "有效 batch=2，P100 16GB 显存极限"),
    ("paged_adamw_8bit", "8-bit 优化器，节省显存"),
]
for i, (k, v) in enumerate(params_data):
    set_cell_text(params_t.rows[i].cells[0], k, bold=True, size=9)
    set_cell_text(params_t.rows[i].cells[1], v, size=9)

# ── 2.4 Training ──
doc.add_heading("2.4 训练", level=2)

doc.add_paragraph("一键训练脚本（已上传到 /data/finetune/run_llamafactory.sh）：")
train_cmd = doc.add_paragraph()
train_cmd.add_run(
    "cd /data/finetune\n"
    "export PATH=/data/envs/qwen-api/bin:$PATH\n"
    "CUDA_VISIBLE_DEVICES=0 llamafactory-cli train train_config.yaml"
).font.size = Pt(9)

doc.add_paragraph("或直接运行：")
train_cmd2 = doc.add_paragraph()
train_cmd2.add_run("bash /data/finetune/run_llamafactory.sh").font.size = Pt(9)

doc.add_paragraph("训练过程关键输出:")
doc.add_paragraph("[INFO] Loading dataset kg_robot_llamafactory.json...")
doc.add_paragraph("Generating train split: 24 examples")
doc.add_paragraph("[INFO] Using torch SDPA for faster training and inference.")
doc.add_paragraph("[INFO] all params: 7,615,616,512")
doc.add_paragraph("可训练参数: ~40M / 7.6B (0.92%)")

doc.add_paragraph()
doc.add_paragraph("训练结果 (2026-05-22):")
result_table = doc.add_table(rows=6, cols=2, style="Table Grid")
result_data = [
    ("训练数据", "24 条（3 原始 + 21 增强）"),
    ("训练时间", "约 6 分钟 (3 epochs)"),
    ("Loss", "2.4 → 0.1"),
    ("LoRA 参数", "40M / 7.6B (0.92%)"),
    ("Checkpoint", "/data/finetune/llamafactory_output/checkpoint-6"),
    ("框架版本", "LLaMA-Factory v0.9.3"),
]
for i, (k, v) in enumerate(result_data):
    set_cell_shading(result_table.rows[i].cells[0], "2F5496")
    set_cell_shading(result_table.rows[i].cells[1], "D6E4F0")
    set_cell_text(result_table.rows[i].cells[0], k, bold=True, size=10, color=(255, 255, 255))
    set_cell_text(result_table.rows[i].cells[1], v, size=10)

# ── 2.5 Merge ──
doc.add_heading("2.5 模型合并", level=2)

doc.add_paragraph(
    "LLaMA-Factory 的 export 命令在 P100 上会 OOM（显存/内存不足）。"
    "因此使用 PEFT 手动合并，直接在 GPU 上操作并保存，不经过 CPU："
)

doc.add_paragraph("合并脚本 (merge_lora.py):")
merge_code = doc.add_paragraph()
merge_code.add_run(
    "import torch\n"
    "from transformers import AutoModelForCausalLM, AutoTokenizer\n"
    "from peft import PeftModel\n\n"
    "adapter = '/data/finetune/llamafactory_output/checkpoint-6'\n"
    "merged = '/data/finetune/output/qwen2.5-7b-kg-robot-merged-v2'\n\n"
    "model = AutoModelForCausalLM.from_pretrained(\n"
    "    '/media/z/data/models/Qwen2.5-7B-Instruct',\n"
    "    torch_dtype=torch.float16,\n"
    "    device_map='cuda:0',\n"
    "    trust_remote_code=True,\n"
    "    low_cpu_mem_usage=True,\n"
    ")\n"
    "model = PeftModel.from_pretrained(model, adapter)\n"
    "model = model.merge_and_unload()\n"
    "model.save_pretrained(merged, safe_serialization=True, max_shard_size='5GB')\n"
    "print(f'Done: {merged}')"
).font.size = Pt(8)

doc.add_paragraph()
doc.add_paragraph("运行:")
doc.add_paragraph("python3 /data/finetune/merge_lora.py")
doc.add_paragraph("输出: /data/finetune/output/qwen2.5-7b-kg-robot-merged-v2/")
doc.add_paragraph(
    "注意事项: 基座模型 /media/z/data/models/Qwen2.5-7B-Instruct 保持不变，"
    "合并后的模型输出到独立路径。"
)

# ── 2.6 Deploy ──
doc.add_heading("2.6 API 部署", level=2)

doc.add_paragraph("步骤 1：更新 API 配置")
deploy_cmd = doc.add_paragraph()
deploy_cmd.add_run(
    "sed -i \"s|^MODEL_PATH = .*|MODEL_PATH = "
    "\"/data/finetune/output/qwen2.5-7b-kg-robot-merged-v2\"|\" /data/qwen_api.py"
).font.size = Pt(9)

doc.add_paragraph("步骤 2：启动 API 服务")
deploy_cmd2 = doc.add_paragraph()
deploy_cmd2.add_run("CUDA_VISIBLE_DEVICES=0 python /data/qwen_api.py").font.size = Pt(9)

doc.add_paragraph("步骤 3：验证")
deploy_cmd3 = doc.add_paragraph()
deploy_cmd3.add_run("curl http://10.117.29.24:5200/health").font.size = Pt(9)
doc.add_paragraph("返回: {\"status\":\"ok\"}")

doc.add_paragraph()
doc.add_paragraph("P100 API 关键配置:")
config_table = doc.add_table(rows=3, cols=2, style="Table Grid")
config_data = [
    ("MAX_MEMORY", '{0: "15GiB"} — 只用 GPU 0，不用双卡配置'),
    ("device_map", '"cuda:0" — 显式指定单卡，不用 "auto"'),
    ("CUDA_VISIBLE_DEVICES", "0 — 环境变量限制单卡'),
]
for i, (k, v) in enumerate(config_data):
    set_cell_shading(config_table.rows[i].cells[0], "2F5496")
    set_cell_shading(config_table.rows[i].cells[1], "D6E4F0")
    set_cell_text(config_table.rows[i].cells[0], k, bold=True, size=10, color=(255, 255, 255))
    set_cell_text(config_table.rows[i].cells[1], v, size=10)

# ── 2.7 Knowledge Graph Integration ──
doc.add_heading("2.7 知识图谱入库", level=2)

doc.add_paragraph("KG 后端启动:")
doc.add_paragraph(
    "cd E:\\Knowledge Graph_robot && "
    "E:\\kg_robot_env\\Scripts\\python.exe -m uvicorn api.app:app --host 0.0.0.0 --port 8000"
)

doc.add_paragraph("文本入库 (batch_upload_text.py):")
doc.add_paragraph(
    "调用 /api/v1/ingest/text → LLMExtractor → "
    "P100 微调模型 (http://10.117.29.24:5200/v1) → Neo4j"
)
doc.add_paragraph("运行: 08_upload_text.bat 或 python batch_upload_text.py")

doc.add_paragraph()
doc.add_paragraph("⚠️ 注意: DWG 文件直接上传走 /ingest/file (DXF 规则引擎)，"
                   "会抽出 CAD 图元 (SW_NOTE_0 等)。"
                   "要用微调模型获取语义知识，必须通过 /ingest/text 上传 .txt 描述。")

# ── 2.8 Full Workflow Diagram ──
doc.add_heading("2.8 完整工作流总览", level=2)

# Color-coded pipeline
pipe = doc.add_table(rows=1, cols=7, style="Table Grid")
pipe.alignment = WD_TABLE_ALIGNMENT.CENTER

pipe_data = [
    ("DWG\n文件", "4472C4"),
    ("txt\n描述", "4472C4"),
    ("增强标注\n+清洗\n+增强", "4472C4"),
    ("LLaMA-Factory\nQLoRA训练", "ED7D31"),
    ("PEFT\n合并模型", "A5A5A5"),
    ("LLM抽取\n→Neo4j", "70AD47"),
    ("/chat\n问答", "FFC000"),
]

for i, (text, color) in enumerate(pipe_data):
    set_cell_shading(pipe.rows[0].cells[i], color)
    set_cell_text(pipe.rows[0].cells[i], text, bold=True, size=9,
                  color=(255, 255, 255), align=WD_ALIGN_PARAGRAPH.CENTER)

# Arrows
arrow_row = pipe.add_row()
for i in range(7):
    set_cell_shading(arrow_row.cells[i], "FFFFFF")
arrow_labels = [(0, "▶", (68, 114, 196)), (1, "▶", (68, 114, 196)),
                (2, "▶", (237, 125, 49)), (3, "▶", (165, 165, 165)),
                (4, "▶", (112, 173, 71)), (5, "▶", (255, 192, 0))]
for i, symbol, color in arrow_labels:
    set_cell_text(arrow_row.cells[i], symbol, size=14, bold=True,
                  color=color, align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_paragraph()

# Detailed 5-stage flow
doc.add_heading("详细阶段流程", level=3)

stages = [
    ("阶段一：数据准备 (Windows)", "D6E4F0",
     [
         "原始 DWG 文件 (6张) → 大模型生成 .txt 自然语言描述",
         "enhanced_label.py: 润色 + 两阶段抽取 + Schema 校验 + 质量评分",
         "clean_labeled.py: 去除 CAD 图元/噪音 + 修正关系 + 规范化实体名",
         "数据增强: 24 条 LLaMA-Factory 格式 (prompt/query/response/system/history)",
         "scp 上传到 P100:/data/finetune/data/",
     ]),
    ("阶段二：QLoRA 训练 (P100)", "FDE8D0",
     [
         "LLaMA-Factory CLI: llamafactory-cli train train_config.yaml",
         "4-bit nf4 量化加载 Qwen2.5-7B → LoRA rank=16",
         "训练: 24条, 3 epochs, batch=1×2, lr=2e-4, fp16",
         "约 6 分钟, Loss 2.4→0.1, 可训练参数 40M/7.6B",
         "输出: checkpoint-6 (adapter_config.json + adapter_model.safetensors)",
     ]),
    ("阶段三：模型合并与部署 (P100)", "E8E8E8",
     [
         "PEFT merge_and_unload(): GPU 直接合并 → 保存 (不经过 CPU 避免 OOM)",
         "合并模型: /data/finetune/output/qwen2.5-7b-kg-robot-merged-v2/",
         "更新 /data/qwen_api.py MODEL_PATH → 启动 API (CUDA_VISIBLE_DEVICES=0)",
         "端点: http://10.117.29.24:5200/v1 → curl /health 验证",
     ]),
    ("阶段四：知识入库 (KG → P100)", "E2EFDA",
     [
         "启动 KG 后端: uvicorn api.app:app --port 8000",
         "batch_upload_text.py: 遍历 .txt → POST /api/v1/ingest/text",
         "LLMExtractor 调用 P100 微调模型 → JSON (entities + relations)",
         "写入 Neo4j: MERGE 节点 + 关系边 (Component/Material/Process 等)",
     ]),
    ("阶段五：问答查询 (浏览器)", "FFF2CC",
     [
         "http://localhost:8000/chat → 输入问题",
         "KG 后端检索 Neo4j → 相关实体和关系作为上下文",
         "LLM 生成自然语言答案 → 前端展示答案 + 知识卡片",
     ]),
]

for title, color, items in stages:
    t = doc.add_table(rows=1, cols=1, style="Table Grid")
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_cell_shading(t.rows[0].cells[0], color)
    cell_text = title + "\n" + "\n".join(f"  {i+1}. {item}" for i, item in enumerate(items))
    set_cell_text(t.rows[0].cells[0], cell_text, size=9)
    doc.add_paragraph()

# Legend
doc.add_heading("图例", level=2)
legend = doc.add_table(rows=1, cols=5, style="Table Grid")
legend.alignment = WD_TABLE_ALIGNMENT.CENTER
legend_colors = ["4472C4", "ED7D31", "A5A5A5", "70AD47", "FFC000"]
legend_labels = ["数据准备 (Windows)", "训练 (P100)", "合并部署 (P100)", "知识入库 (KG)", "问答 (浏览器)"]
for i, (c, l) in enumerate(zip(legend_colors, legend_labels)):
    set_cell_shading(legend.rows[0].cells[i], c)
    set_cell_text(legend.rows[0].cells[i], l, bold=True, size=9,
                  color=(255, 255, 255), align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_page_break()

# ═══════════════════════════════════════
# PART 3: Key Issues and Fixes
# ═══════════════════════════════════════
doc.add_heading("第三部分：关键问题与解决方案", level=1)

fixes = [
    ("测试 Prompt 不匹配",
     "训练时使用 ~800 字符完整 instruction（实体/关系类型列表 + JSON 格式规则 + 5条抽取规则），"
     "测试时只用了 ~50 字符通用消息 → 模型输出自然语言而非 JSON。"
     "修复: 测试推理时使用与训练完全一致的 instruction。"),

    ("P100 GPU 配置错误",
     "MAX_MEMORY={0:'12GiB',1:'12GiB'} + device_map='auto' → CUDA error: invalid device ordinal。"
     "修复: MAX_MEMORY={0:'15GiB'} + device_map='cuda:0'，只配置实际使用的单卡。"),

    ("数据增强死循环",
     "向 samples 列表追加增强样本时同时在遍历该列表 → 无限循环。"
     "修复: 先快照原始样本遍历，新样本收集到独立列表，结束后再合并。"),

    ("DXF 规则引擎 vs LLM 抽取",
     "/ingest/file (DWG) → DXF 规则引擎 → 抽出 CAD 图元 (SW_NOTE_0, SW_CENTERMARKSYMBOL_0)。"
     "/ingest/text (.txt) → LLMExtractor → 微调模型 → 语义实体。"
     "修复: 要使用微调模型必须走文本入库路径，DWG 直接上传走规则引擎。"),

    ("LLaMA-Factory 导出 OOM",
     "llamafactory-cli export 默认 bf16 加载模型 + CPU merge → P100 内存/显存不足被杀。"
     "修复: 使用 PEFT merge_and_unload() 直接在 GPU 合并并保存，不经过 CPU。"),

    ("P100 不支持 bf16",
     "bf16=true 会在训练时崩溃。"
     "修复: 配置中设置 fp16=true, bf16=false。P100 会自动使用 float16。"),
]

for t, desc in fixes:
    doc.add_heading(t, level=2)
    doc.add_paragraph(desc)

doc.add_page_break()

# ═══════════════════════════════════════
# PART 4: File Reference
# ═══════════════════════════════════════
doc.add_heading("第四部分：关键文件索引", level=1)

files_table = doc.add_table(rows=17, cols=3, style="Table Grid")
files_table.alignment = WD_TABLE_ALIGNMENT.CENTER

# Header
for i, h in enumerate(["位置", "文件", "说明"]):
    set_cell_shading(files_table.rows[0].cells[i], "2F5496")
    set_cell_text(files_table.rows[0].cells[i], h, bold=True, size=9, color=(255, 255, 255))

files_data = [
    ("Windows", "E:\\...\\finetune\\enhanced_label.py", "增强标注：润色+抽取+校验+增强"),
    ("Windows", "E:\\...\\finetune\\clean_labeled.py", "自动清洗：去CAD图元+修正关系+规范化"),
    ("Windows", "E:\\...\\finetune\\batch_upload_text.py", "文本入库：遍历.txt→LLM抽取→Neo4j"),
    ("Windows", "E:\\...\\finetune\\data\\", "训练数据目录"),
    ("Windows", "C:\\Users\\Knightz\\Desktop\\train_dwg\\", "原始DWG+.txt+批处理文件"),
    ("P100", "/data/finetune/train_config.yaml", "LLaMA-Factory 训练配置"),
    ("P100", "/data/finetune/data/dataset_info.json", "数据集注册"),
    ("P100", "/data/finetune/data/kg_robot_llamafactory.json", "训练数据 (24条)"),
    ("P100", "/data/finetune/run_llamafactory.sh", "一键训练+导出+部署脚本"),
    ("P100", "/data/finetune/merge_lora.py", "PEFT 合并脚本 (解决OOM)"),
    ("P100", "/data/finetune/llamafactory_output/", "训练输出 (checkpoint)"),
    ("P100", "/data/finetune/output/qwen2.5-7b-kg-robot-merged-v2/", "合并后完整模型"),
    ("P100", "/data/qwen_api.py", "API 服务脚本"),
    ("P100 (只读)", "/media/z/data/models/Qwen2.5-7B-Instruct", "基座模型 (不修改)"),
    ("P100", "/data/envs/qwen-api/", "Python 环境 (含 LLaMA-Factory)"),
    ("Web", "http://10.117.29.24:5200/v1", "P100 API 端点"),
]

for i, (loc, path, desc) in enumerate(files_data):
    set_cell_text(files_table.rows[i+1].cells[0], loc, size=8, bold=True)
    set_cell_text(files_table.rows[i+1].cells[1], path, size=8)
    set_cell_text(files_table.rows[i+1].cells[2], desc, size=8)

doc.add_paragraph()

# ── Save ──
doc.save(str(OUTPUT))
print(f"Saved: {OUTPUT}")
