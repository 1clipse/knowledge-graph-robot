from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os

doc = Document()

# -- Styles --
style = doc.styles['Normal']
font = style.font
font.name = 'Microsoft YaHei'
font.size = Pt(11)

# -- Title --
title = doc.add_heading('工业机器人知识图谱 v2 版本更新说明', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('发布日期：2026-05-25 | 项目路径：E:\\Knowledge Graph_robot\\')
doc.add_paragraph('')

# -- Overview --
doc.add_heading('一、版本概述', 1)
doc.add_paragraph(
    'v2 版本在 v1 的基础上进行了 7 项重大改进，涵盖安全性、检索质量、数据可信度、'
    '数据治理、用户体验和时间维度支持。核心目标是将系统从"能用"提升到"好用"，'
    '使之接近业界主流知识图谱产品的水平。'
)

# -- Feature Details --
doc.add_heading('二、功能详解', 1)

features = [
    {
        'title': 'P0-1 安全修复',
        'desc': (
            'config/default.yaml 中的 Neo4j AuraDB 明文密码已移除，替换为占位符 "changeme"。'
            '真实密码仅保存在 config/.env 中，该文件已在 .gitignore 中排除，不会被版本控制追踪。'
            'P100 部署时通过系统环境变量注入密码，进一步降低泄露风险。'
        ),
    },
    {
        'title': 'P0-2 多跳推理',
        'desc': (
            '问答模块的核心升级。v1 版本的 /ask 仅做关键词搜索 → 1-hop 邻居展开 → LLM 回答，'
            '无法回答需要推理链的问题（如"FANUC M-20iA 的减速器供应商是谁？"需要 Robot→Reducer→Manufacturer 两跳）。\n\n'
            'v2 实现方案：\n'
            '1. graph/query.py 新增 multi_hop_paths() 方法，从种子实体出发沿关系路径 BFS 遍历 2-3 跳\n'
            '2. api/routes/ask.py 重构 _build_context()，构建完整推理链：\n'
            '   路径1: [Robot]M-20iA --[uses_reducer]--> [Reducer]RV-40E --[supplies_component]--> [Manufacturer]Nabtesco\n'
            '3. LLM 能同时看到多跳关系链，进行跨跳推理\n'
            '4. 流式端点 /ask/stream 同步支持'
        ),
    },
    {
        'title': 'P1-1 向量语义检索',
        'desc': (
            'v1 版本仅使用 Neo4j 全文索引做关键词匹配，无法处理语义相近的表达'
            '（如"焊接机械臂"搜不到"弧焊机器人"）。\n\n'
            'v2 实现方案：\n'
            '1. 新增 graph/embeddings.py，使用 BAAI/bge-small-zh-v1.5 模型（384 维向量）\n'
            '2. 摄入时自动为每个实体生成语义向量，存入 _embedding 属性\n'
            '3. 查询时执行混合检索：向量余弦相似度 + Neo4j 全文索引，结果去重合并排序\n'
            '4. 如果 sentence-transformers 未安装，自动降级为纯关键词检索（无依赖阻塞）\n\n'
            '技术细节：\n'
            '- 中文优化：BGE-small-zh 专为中文语义检索设计\n'
            '- 轻量级：模型仅 130MB，CPU 可运行，GPU 加速更快\n'
            '- 归一化向量：cosine_similarity = dot_product，无需额外计算'
        ),
    },
    {
        'title': 'P1-2 知识溯源（Confidence & Source）',
        'desc': (
            'v1 版本所有入库的三元组无置信度标注，无法区分"从 CAD 图纸精确提取的尺寸参数"'
            '和"LLM 推测的可能关系"。\n\n'
            'v2 实现方案：\n'
            '1. ExtractedEntity / ExtractedRelation 模型新增字段：\n'
            '   - source: 来源标识（文件名/URL/"text"）\n'
            '   - source_text: 原始文本片段\n'
            '   - confidence: 置信度（0-1 浮点数）\n'
            '2. 不同抽取路径的置信度：\n'
            '   - LLM 抽取: 0.7（正确率高但偶有幻觉）\n'
            '   - 正则规则: 0.95（精确匹配但覆盖有限）\n'
            '   - CSV 结构化映射: 1.0（人工整理的数据）\n'
            '   - DXF CAD 解析: 0.95（结构化图元提取）\n'
            '3. api/routes/ingest.py 新增 _stamp_source() 函数统一标注来源\n'
            '4. Neo4j 写入时自动附带 _source / _confidence 属性\n'
            '5. 前端节点详情面板展示置信度和来源文件'
        ),
    },
    {
        'title': 'P1-3 图谱质量面板',
        'desc': (
            '新增 GET /api/v1/quality 端点（api/routes/quality.py），提供数据质量报告。\n\n'
            '返回内容：\n'
            '- quality_score: 0-100 综合评分（100 减去问题数×2）\n'
            '- orphan_nodes: 无任何关系的孤立节点（排除 IngestLog）\n'
            '- missing_key_props: 缺失关键属性的实体（如 Robot 缺 payload，Reducer 缺 reducer_type）\n'
            '- potential_duplicates: 同名异类型的潜在重复实体\n'
            '- low_confidence_facts: 置信度低于 0.5 的实体/关系\n'
            '- schema_violations: 空关系类型、缺失名称等 schema 违规\n\n'
            '界面入口：前端图谱统计栏右侧"质量报告"按钮。'
        ),
    },
    {
        'title': 'P2-1 对比分析',
        'desc': (
            '新增 POST /api/v1/compare 端点，支持对任意两个实体生成结构化对比报告。\n\n'
            '实现流程：\n'
            '1. 搜索两个实体（支持模糊匹配）\n'
            '2. 获取各自属性、邻居关系、共同关系类型\n'
            '3. 构建对比上下文，调用 LLM 生成 Markdown 格式报告\n'
            '4. 前端渲染对比结果\n\n'
            '用法示例：\n'
            '  输入: "FANUC M-20iA" vs "ABB IRB 6700"\n'
            '  输出: 逐项对比负载/臂展/精度/轴数/应用场景，标注各自的优势和差异\n\n'
            '界面入口：前端侧边栏"对比分析"面板。'
        ),
    },
    {
        'title': 'P2-2 时序知识图谱',
        'desc': (
            'v1 版本所有事实无时间维度，无法表达"该型号从何时开始生产、何时停产"。\n\n'
            'v2 实现方案：\n'
            '1. ExtractedEntity / ExtractedRelation 新增 valid_from / valid_to 字段\n'
            '2. LLM 抽取 prompt 增加时间提取 few-shot 示例\n'
            '   例: "FANUC于2010年推出M-20iA，2024年宣布停产"\n'
            '   → relation: manufactures, valid_from: 2010, valid_to: 2024\n'
            '3. 写入 Neo4j 时自动存储 valid_from / valid_to 属性\n'
            '4. 新增 API 端点：\n'
            '   - GET /api/v1/timeline: 全图时间线\n'
            '   - GET /api/v1/timeline/{name}: 单个实体的历史变迁\n'
            '5. 前端节点详情展示时间范围'
        ),
    },
]

for f in features:
    doc.add_heading(f['title'], 2)
    doc.add_paragraph(f['desc'])

# -- Technical Architecture --
doc.add_heading('三、技术实现架构', 1)
doc.add_paragraph('')

doc.add_heading('3.1 新增文件', 2)
doc.add_paragraph('api/routes/quality.py — 质量报告端点')
doc.add_paragraph('graph/embeddings.py — BGE 向量生成 + 余弦相似度计算')

doc.add_heading('3.2 修改文件', 2)
changes = [
    'config/default.yaml — 移除明文密码',
    'extractors/llm_extractor.py — 模型字段扩展 + prompt 优化 + 置信度/时序赋值',
    'extractors/rule_extractor.py — 规则抽取加 0.95 置信度',
    'graph/query.py — 新增 multi_hop_paths() 多跳遍历',
    'api/routes/ask.py — 完全重构：_hybrid_search() + _build_context() 多跳 + /compare 端点',
    'api/routes/ingest.py — _stamp_source() + _write_to_graph() 溯源/置信度/embedding/时序',
    'api/routes/query.py — 新增 /timeline + /timeline/{name} 端点',
    'api/app.py — 注册 quality 路由',
    'ui/index.html — 对比分析面板 + 质量报告按钮 + 节点详情溯源展示',
    'requirements.txt — 新增 sentence-transformers, ezdxf',
]
for c in changes:
    doc.add_paragraph(c, style='List Bullet')

doc.add_heading('3.3 检索链路升级', 2)
doc.add_paragraph(
    'v1 链路: 用户问题 → 全文关键词索引 → 1-hop 邻居 → 拼文本 → LLM\n\n'
    'v2 链路: 用户问题 → 向量相似度 + 关键词混合检索 → 去重合并 → '
    '多跳路径遍历(2-3跳) → 推理链构建 → LLM'
)

doc.add_heading('3.4 新增依赖', 2)
doc.add_paragraph('sentence-transformers >= 2.2.0 — 语义向量模型 (BGE-small-zh-v1.5, ~130MB)')
doc.add_paragraph('ezdxf >= 1.0.0 — DXF 图纸解析（P100 Linux 端需要）')
doc.add_paragraph('安装命令: pip install sentence-transformers ezdxf')

# -- API Reference --
doc.add_heading('四、API 端点速查', 1)

table = doc.add_table(rows=1, cols=3)
table.style = 'Light Grid Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = '端点'
hdr_cells[1].text = '方法'
hdr_cells[2].text = '功能'

endpoints = [
    ('/api/v1/quality', 'GET', '图谱质量报告'),
    ('/api/v1/compare', 'POST', '两实体对比分析'),
    ('/api/v1/timeline', 'GET', '全图时间线'),
    ('/api/v1/timeline/{name}', 'GET', '单实体历史'),
    ('/api/v1/ask', 'POST', '多跳推理问答（升级）'),
    ('/api/v1/ask/stream', 'POST', '流式问答（升级）'),
]

for ep, method, desc in endpoints:
    row_cells = table.add_row().cells
    row_cells[0].text = ep
    row_cells[1].text = method
    row_cells[2].text = desc

# -- Deployment --
doc.add_heading('五、部署注意事项', 1)
doc.add_paragraph(
    '1. 本地测试：Windows 上 cd E:\\Knowledge Graph_robot && '
    'pip install sentence-transformers ezdxf && '
    'set PYTHONPATH=E:\\Knowledge Graph_robot && uvicorn api.app:app --host 0.0.0.0 --port 8100\n\n'
    '2. P100 后端迁移（规划中）：\n'
    '   - conda activate kg-robot (/data/envs/kg-robot)\n'
    '   - pip install sentence-transformers ezdxf\n'
    '   - 修改 config/.env: LLM_BASE_URL=http://127.0.0.1:5200/v1\n'
    '   - uvicorn api.app:app --host 0.0.0.0 --port 8100\n\n'
    '3. 如果 sentence-transformers 未安装，向量检索自动降级为纯关键词检索，不影响系统使用\n\n'
    '4. DWG 图纸摄入需要 ODA File Converter（仅 Windows 端），后续计划 Windows 转换 → P100 摄入'
)

# -- Save --
desktop = os.path.expanduser("~/Desktop")
filepath = os.path.join(desktop, "KG_Robot_v2_更新说明_2026-05-25.docx")
doc.save(filepath)
print(f"Saved: {filepath}")
