"""Generate comprehensive deployment documentation as .docx"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os

doc = Document()

# --- Style configuration ---
style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(10.5)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.25

# --- Title ---
title = doc.add_heading('工业机器人知识图谱系统', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_run = subtitle.add_run('部署与运维手册')
subtitle_run.font.size = Pt(14)
subtitle_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_p.add_run('2026-05-20').font.size = Pt(10)

doc.add_page_break()

# ============================================================
# 1. ARCHITECTURE
# ============================================================
doc.add_heading('一、系统架构总览', level=1)

doc.add_paragraph(
    '系统由四个核心组件构成，分布式部署在 Windows 本地、Linux 服务器和 Neo4j 云端：'
)

# Architecture table
table = doc.add_table(rows=5, cols=4)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
hdr[0].text = '组件'
hdr[1].text = '位置'
hdr[2].text = '端口'
hdr[3].text = '技术栈'

data = [
    ['前端 UI', 'Windows 本机', '8100', 'HTML/CSS/JS + D3.js'],
    ['后端 API', 'Windows 本机', '8100', 'FastAPI + Uvicorn (Python 3.11)'],
    ['LLM 推理', 'Linux 10.117.29.24', '5200', 'Qwen2.5-7B + Transformers + 2x P100'],
    ['图数据库', 'Neo4j AuraDB 云端', '7687', 'b0daa589.databases.neo4j.io'],
]
for i, row_data in enumerate(data):
    for j, val in enumerate(row_data):
        table.rows[i+1].cells[j].text = val

doc.add_paragraph()
doc.add_paragraph(
    '数据流向：前端上传文件 → API 接收 → 文件解析器提取文本 → LLM/规则引擎抽取实体关系 → 写入 Neo4j → 前端图谱可视化'
).italic = True

# ============================================================
# 2. NEO4J
# ============================================================
doc.add_heading('二、Neo4j 图数据库', level=1)

doc.add_heading('2.1 连接信息', level=2)
doc.add_paragraph('使用 Neo4j AuraDB 免费云实例，无需本地安装。')
table2 = doc.add_table(rows=5, cols=2)
table2.style = 'Light Grid Accent 1'
info = [
    ['URI', 'neo4j://b0daa589.databases.neo4j.io:7687'],
    ['用户名', 'b0daa589'],
    ['数据库名', 'b0daa589'],
    ['配置文件', r'E:\Knowledge Graph_robot\config\.env'],
]
for i, (k, v) in enumerate(info):
    table2.rows[i].cells[0].text = k
    table2.rows[i].cells[1].text = v
table2.rows[4].cells[0].text = 'Web 控制台'
table2.rows[4].cells[1].text = 'https://console.neo4j.io'

doc.add_heading('2.2 知识图谱本体 (Schema)', level=2)
doc.add_paragraph('定义在 schema/industrial_robot.yaml，包含以下实体类型：')

labels = [
    'Robot（机器人）、Manufacturer（制造商）、Component（零部件）、Reducer（减速器）、'
    'ServoMotor（伺服电机）、Controller（控制器）、Sensor（传感器）、'
    'ApplicationScenario（应用场景）、Process（工艺）、EndEffector（末端执行器）、'
    'Standard（标准规范）、Material（加工材料）、Software（软件系统）'
]
doc.add_paragraph(', '.join(labels))

doc.add_heading('2.3 验证数据库连接', level=2)
p = doc.add_paragraph()
p.style = doc.styles['List Bullet']
p.add_run('健康检查：').bold = True
p.add_run(' curl http://localhost:8100/health')

p2 = doc.add_paragraph()
p2.style = doc.styles['List Bullet']
p2.add_run('查看统计：').bold = True
p2.add_run(' curl http://localhost:8100/api/v1/query/stats')

# ============================================================
# 3. WINDOWS BACKEND
# ============================================================
doc.add_heading('三、Windows 后端 API', level=1)

doc.add_heading('3.1 项目结构', level=2)
code = doc.add_paragraph()
code.style = doc.styles['Normal']
code_run = code.add_run(
    'E:\\Knowledge Graph_robot\\\n'
    '├── api\\                  # FastAPI 应用\n'
    '│   ├── app.py            # 主入口\n'
    '│   ├── deps.py           # 依赖注入（Neo4j客户端、Schema管理器）\n'
    '│   └── routes\\           # API 路由\n'
    '│       ├── ingest.py     # 数据摄入（文本/文件/URL）\n'
    '│       ├── ask.py        # 知识问答（RAG 流式）\n'
    '│       ├── query.py      # 图查询（邻居、统计、最短路径）\n'
    '│       └── subgraph.py   # 子图检索（可视化用）\n'
    '├── extractors\\           # 实体关系抽取器\n'
    '│   ├── llm_extractor.py # LLM 提取（调用远程模型）\n'
    '│   ├── llm_utils.py     # LLM 调用工具\n'
    '│   ├── rule_extractor.py# 规则提取（正则匹配回退）\n'
    '│   └── structured_mapper.py # CSV 结构化映射\n'
    '├── loaders\\              # 文件加载器\n'
    '│   ├── dxf_loader.py    # DXF 解析（ezdxf）\n'
    '│   ├── step_loader.py   # STEP/IGES 解析\n'
    '│   ├── pdf_loader.py    # PDF 解析\n'
    '│   ├── csv_loader.py    # CSV 解析\n'
    '│   └── web_loader.py    # 网页抓取\n'
    '├── graph\\                # Neo4j 交互层\n'
    '│   ├── client.py        # Neo4j 客户端（CRUD）\n'
    '│   ├── query.py         # 图查询（全文搜索、子图、路径）\n'
    '│   └── schema_manager.py# Schema 管理（约束、索引）\n'
    '├── schema\\               # 知识图谱本体\n'
    '│   ├── industrial_robot.yaml  # 实体/关系定义\n'
    '│   └── loader.py        # Schema 加载器\n'
    '├── config\\               # 配置\n'
    '│   ├── .env              # 环境变量（密钥）\n'
    '│   ├── default.yaml     # 默认配置\n'
    '│   └── settings.py      # Pydantic Settings\n'
    '├── ui\\                   # 前端\n'
    '│   └── index.html       # 单页应用\n'
    '└── pipeline\\             # 数据处理管道\n'
)
code_run.font.name = 'Cascadia Code'
code_run.font.size = Pt(8)

doc.add_heading('3.2 依赖安装', level=2)
doc.add_paragraph('Python 3.11，关键依赖：')
deps = '''fastapi uvicorn[standard] neo4j openai ezdxf pydantic-settings
loguru pyyaml tenacity httpx rapidfuzz pypdf langchain-text-splitters
python-docx'''
p = doc.add_paragraph(deps)
p.style = doc.styles['Normal']

doc.add_heading('3.3 启动命令', level=2)
doc.add_paragraph('在 CMD 或 PowerShell 中执行：')
startup = doc.add_paragraph()
startup_run = startup.add_run(
    'cd E:\\Knowledge Graph_robot\n'
    'set PYTHONPATH=E:\\Knowledge Graph_robot\n'
    'uvicorn api.app:app --host 0.0.0.0 --port 8100 --reload'
)
startup_run.font.name = 'Cascadia Code'
startup_run.font.size = Pt(9)

doc.add_paragraph('注意：必须在项目目录下启动，否则 --reload 监控路径错误导致代码改动不生效。')

doc.add_heading('3.4 关键配置 (.env)', level=2)
dotenv = doc.add_paragraph()
dotenv_run = dotenv.add_run(
    '# Neo4j\n'
    'NEO4J_URI=neo4j://b0daa589.databases.neo4j.io:7687\n'
    'NEO4J_USERNAME=b0daa589\n'
    'NEO4J_PASSWORD=***\n'
    'NEO4J_DATABASE=b0daa589\n\n'
    '# LLM\n'
    'LLM_BASE_URL=http://10.117.29.24:5200/v1\n'
    'LLM_API_KEY=local\n'
    'LLM_MODEL=qwen2.5-7b\n'
    'LLM_TEMPERATURE=0.1\n'
    'LLM_MAX_TOKENS=4096\n\n'
    '# Logging\n'
    'LOG_LEVEL=INFO'
)
dotenv_run.font.name = 'Cascadia Code'
dotenv_run.font.size = Pt(8)

doc.add_heading('3.5 数据摄入 API', level=2)
doc.add_paragraph('支持的文件格式：.pdf .csv .txt .md .dwg .dxf .step .stp .igs .iges')

doc.add_paragraph('DWG 文件处理流程：', style='List Bullet')
items = [
    '上传 DWG → ODA File Converter 转为 DXF',
    'DXFLoader (ezdxf) 解析结构化数据：图层、图块、文字标注、尺寸、属性',
    '_dxf_meta_to_entities() 直接映射为 Component 实体 + contains 关系',
    '同时截断文本(≤3500字) 发送给 LLM 做语义提取补充',
    'LLM 失败时回退到规则引擎（正则匹配）',
    '合并直接映射 + LLM/规则结果 → 写入 Neo4j',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet 2')

# ============================================================
# 4. ODA FILE CONVERTER
# ============================================================
doc.add_heading('四、ODA File Converter (DWG→DXF)', level=1)

doc.add_paragraph('ODA File Converter 是第三方命令行工具，将 DWG 转为 DXF。')
doc.add_paragraph(f'安装路径：E:\\ODA\\ODAFileConverter.exe')
doc.add_paragraph(f'关键要求：exe 同级目录下必须有 platforms/qwindows.dll（Qt6 平台插件）')
doc.add_paragraph()

doc.add_paragraph('CLI 格式：')
p = doc.add_paragraph()
p_run = p.add_run(
    'ODAFileConverter.exe "输入目录" "输出目录" "ACAD2018" "DXF" "0" "0"'
)
p_run.font.name = 'Cascadia Code'
p_run.font.size = Pt(9)

doc.add_paragraph('注意事项：')
items2 = [
    '输入必须是文件夹路径（不是文件），代码会在临时目录复制 DWG 后调用 ODA',
    'subprocess.CREATE_NO_WINDOW (0x08000000) 阻止 Qt GUI 弹窗',
    '超时限制 120 秒',
    '转换失败时回退到二进制字符串扫描（提取可读 ASCII/UTF-8 文本）',
]
for item in items2:
    doc.add_paragraph(item, style='List Bullet')

# ============================================================
# 5. FRONTEND
# ============================================================
doc.add_heading('五、前端界面', level=1)

doc.add_heading('5.1 访问方式', level=2)
doc.add_paragraph('启动后端后浏览器访问：http://localhost:8100')
doc.add_paragraph('单页应用，左侧边栏切换两个视图：')

table3 = doc.add_table(rows=3, cols=2)
table3.style = 'Light Grid Accent 1'
table3.rows[0].cells[0].text = '视图'
table3.rows[0].cells[1].text = '功能'
table3.rows[1].cells[0].text = '图谱浏览'
table3.rows[1].cells[1].text = 'D3.js 力导向图、实体搜索、路径查询、节点详情、数据录入'
table3.rows[2].cells[0].text = '知识问答'
table3.rows[2].cells[1].text = '基于知识图谱的 RAG 流式对话'

doc.add_heading('5.2 数据录入操作', level=2)
ops = [
    '文本录入：在左侧「数据录入」面板输入文本，点"提交录入"',
    '文件上传：点"上传文件"，选择 DWG/DXF/PDF/CSV 等文件',
    '验证结果：看状态栏是否显示"录入成功"，以及实体/关系数量',
    '查看数据：在「图谱浏览」搜索实体名，或全貌总览按钮',
]
for op in ops:
    doc.add_paragraph(op, style='List Bullet')

# ============================================================
# 6. LINUX LLM SERVER
# ============================================================
doc.add_heading('六、Linux LLM 推理服务器', level=1)

doc.add_heading('6.1 硬件配置', level=2)
table4 = doc.add_table(rows=4, cols=2)
table4.style = 'Light Grid Accent 1'
hw = [
    ['服务器', 'z@z-Z590-D (IP: 10.117.29.24)'],
    ['GPU', '2x Tesla P100-SXM2-16GB (15.89 GiB each)'],
    ['模型', 'Qwen2.5-7B-Instruct (~14GB, float16)'],
]
for i, (k, v) in enumerate(hw):
    table4.rows[i].cells[0].text = k
    table4.rows[i].cells[1].text = v
table4.rows[3].cells[0].text = '架构限制'
table4.rows[3].cells[1].text = 'P100 = Pascal SM 6.0，不支持 vLLM (需 Volta SM 7.0+)，只能用 Transformers 推理'

doc.add_heading('6.2 部署文件', level=2)
doc.add_paragraph('/data/qwen_api.py — FastAPI 封装的 OpenAI 兼容 API')
doc.add_paragraph('/data/models/Qwen2.5-7B-Instruct — 模型权重')
doc.add_paragraph('/data/envs/qwen-api — Python 虚拟环境')

doc.add_heading('6.3 关键配置参数', level=2)
params = doc.add_paragraph()
params_run = params.add_run(
    'import torch\n'
    'MODEL_PATH = "/data/models/Qwen2.5-7B-Instruct"\n'
    'DTYPE = torch.float16\n'
    'MAX_MEMORY = {0: "11GiB", 1: "11GiB"}  # 每卡留约5GB给KV cache\n'
    '\n'
    'model = AutoModelForCausalLM.from_pretrained(\n'
    '    MODEL_PATH,\n'
    '    torch_dtype=DTYPE,\n'
    '    device_map="auto",      # 自动分配到两张GPU\n'
    '    max_memory=MAX_MEMORY,  # 限制每卡显存使用\n'
    '    trust_remote_code=True,\n'
    ')'
)
params_run.font.name = 'Cascadia Code'
params_run.font.size = Pt(8)

doc.add_heading('6.4 显存分配原理', level=2)
doc.add_paragraph(
    'Qwen2.5-7B float16 ~14GB 总权重。device_map="auto" 将 28 层 Transformer 均分到两张卡，'
    '每卡 ~7GB 权重 + KV cache。max_memory="11GiB" 限制后每卡有 ~4GB 给缓存，'
    '可支持约 8000 tokens 上下文。若设 "14GiB" 则每卡仅剩 ~2GB，长 prompt 即 OOM。'
)

doc.add_heading('6.5 启动命令', level=2)
startup2 = doc.add_paragraph()
startup2_run = startup2.add_run(
    '# 激活环境\n'
    'conda activate qwen-api\n\n'
    '# 前台启动（看日志）\n'
    'CUDA_VISIBLE_DEVICES=0,1 python /data/qwen_api.py\n\n'
    '# 后台启动\n'
    'CUDA_VISIBLE_DEVICES=0,1 nohup python /data/qwen_api.py > /data/qwen.log 2>&1 &\n\n'
    '# 查看状态\n'
    'nvidia-smi\n'
    'curl -s http://10.117.29.24:5200/health\n\n'
    '# 停服\n'
    'pkill -f qwen_api'
)
startup2_run.font.name = 'Cascadia Code'
startup2_run.font.size = Pt(9)

doc.add_heading('6.6 测试 API', level=2)
test = doc.add_paragraph()
test_run = test.add_run(
    'curl -s http://10.117.29.24:5200/v1/chat/completions \\\n'
    '  -H "Content-Type: application/json" \\\n'
    '  -d \'{"model":"qwen2.5-7b","messages":[{"role":"user","content":"你好"}],"max_tokens":50}\''
)
test_run.font.name = 'Cascadia Code'
test_run.font.size = Pt(9)

# ============================================================
# 7. TROUBLESHOOTING
# ============================================================
doc.add_heading('七、常见问题与排查', level=1)

problems = [
    ('DWG 上传后 0 实体 0 关系',
     '1. 检查 uvicorn 终端日志：是否有 "DXF loaded" 和 "direct entities"\n'
     '2. 如果日志显示旧格式：uvicorn --reload 未生效，检查是否在项目目录启动\n'
     '3. 如果 LLM 报 "peer closed connection"：LLM 服务器 OOM，检查 nvidia-smi\n'
     '4. 如果规则引擎也是 0：DXF 文本格式不匹配规则模式（这是已知问题，直接映射层已解决）'),
    ('LLM 服务器 CUDA OOM',
     '1. ssh 到 10.117.29.24\n'
     '2. nvidia-smi 看显存占用，如果接近 100% 说明 KV cache 不足\n'
     '3. 检查 max_memory 配置：grep MAX_MEMORY /data/qwen_api.py\n'
     '4. 修改为 {0: "11GiB", 1: "11GiB"}，重启\n'
     '5. 如果仍然 OOM，可能需要换更小的模型或增加 GPU'),
    ('ODA 转换失败 (rc=3221226505)',
     '1. 检查 E:\\ODA\\platforms\\qwindows.dll 是否存在\n'
     '2. 不要设置 QT_QPA_PLATFORM 环境变量\n'
     '3. 确认使用 subprocess.CREATE_NO_WINDOW 标志'),
    ('前端搜索无结果',
     '1. 全文索引对下划线不分词，"SW_CENTERMARKS" 搜不到 "SW_CENTERMARKS_33"\n'
     '2. 已修复：全文搜不到自动降级 CONTAINS 子串匹配\n'
     '3. 如果仍然不行，重启 uvicorn 确保代码生效'),
    ('uvicorn 找不到 api 模块',
     '必须设置 PYTHONPATH：set PYTHONPATH=E:\\Knowledge Graph_robot'),
    ('DWG 文件上传"正在上传"一直不返回',
     '1. ODA 转换大文件需要时间（超时 120s）\n'
     '2. LLM 调用有 120s 超时\n'
     '3. 看 uvicorn 终端日志确认卡在哪一步'),
]
for prob, sol in problems:
    doc.add_heading(prob, level=2)
    doc.add_paragraph(sol)

# ============================================================
# 8. FULL STARTUP CHECKLIST
# ============================================================
doc.add_heading('八、完整启动检查清单', level=1)

checklist = [
    '□ Linux 服务器：conda activate qwen-api',
    '□ Linux 服务器：nvidia-smi 确认两张 P100 空闲',
    '□ Linux 服务器：CUDA_VISIBLE_DEVICES=0,1 python /data/qwen_api.py',
    '□ Linux 服务器：curl http://10.117.29.24:5200/health 返回 {"status":"ok"}',
    '□ Windows：cd E:\\Knowledge Graph_robot',
    '□ Windows：set PYTHONPATH=E:\\Knowledge Graph_robot',
    '□ Windows：uvicorn api.app:app --host 0.0.0.0 --port 8100 --reload',
    '□ Windows：浏览器访问 http://localhost:8100',
    '□ Windows：左侧「图谱统计」显示节点/关系数量',
    '□ 测试：左侧「知识问答」，提问 "FANUC 有哪些机器人"',
    '□ 测试：左侧「数据录入」→ 上传文件 → 选择 DWG → 确认返回实体数 > 0',
]
for item in checklist:
    doc.add_paragraph(item)

# ============================================================
# SAVE
# ============================================================
output_path = os.path.expanduser(r'~\Desktop\知识图谱系统部署手册.docx')
doc.save(output_path)
print(f'Document saved to: {output_path}')
