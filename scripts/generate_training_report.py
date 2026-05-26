#!/usr/bin/env python3
"""Generate .docx report for QLoRA fine-tuning session 2026-05-22"""
from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# ── Styles ──
style = doc.styles["Normal"]
font = style.font
font.name = "Microsoft YaHei"
font.size = Pt(11)

# ── Title ──
title = doc.add_heading("Qwen2.5-7B QLoRA Fine-tuning Report", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph(f"Date: 2026-05-22")
doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")

# ── 1. Overview ──
doc.add_heading("1. Overview", level=1)
doc.add_paragraph(
    "Fine-tuned Qwen2.5-7B-Instruct on industrial robot DWG drawings for "
    "knowledge graph entity and relation extraction. Used QLoRA (4-bit NF4 "
    "quantization + Low-Rank Adaptation) to train efficiently on a P100 server."
)

doc.add_heading("1.1 Hardware", level=2)
doc.add_paragraph("GPU: 2x Tesla P100-SXM2-16GB (only GPU 0 used for training)")
doc.add_paragraph("Server: x86_64, ZeroTier IP 10.117.29.24")
doc.add_paragraph("Training time: ~6.5 minutes for 3 epochs")

doc.add_heading("1.2 Training Data", level=2)
doc.add_paragraph("Source: 6 DWG files from C:\\Users\\Knightz\\Desktop\\train_dwg")
doc.add_paragraph("Final dataset: 27 samples (3 handcrafted + 24 augmented)")
doc.add_paragraph("Entity types: 13 (Robot, Manufacturer, Component, Reducer, etc.)")
doc.add_paragraph("Relation types: 10 (has_part, made_of, performs_process, etc.)")

# ── 2. Pipeline ──
doc.add_heading("2. Pipeline Steps", level=1)

doc.add_heading("Step 1: Enhanced Labeling (enhanced_label.py)", level=2)
doc.add_paragraph("Polish DWG descriptions - remove CAD junk (%%c, %%p), rewrite as natural language")
doc.add_paragraph("Two-stage extraction: entities first, then relations based on entities")
doc.add_paragraph("Schema validation against industrial robot ontology")
doc.add_paragraph("Quality scoring: good / ok / review")
doc.add_paragraph("Output: enhanced_labeled.json + LLaMA-Factory format")

doc.add_heading("Step 2: Auto Cleaning (clean_labeled.py)", level=2)
doc.add_paragraph("Remove CAD primitives: circles, arcs, lines, blocks, center marks")
doc.add_paragraph("Remove noise entities: standard materials, dimension info, annotations")
doc.add_paragraph("Fix relation direction: contains (container->contained), performs_process")
doc.add_paragraph("Deduplicate bidirectional relations")
doc.add_paragraph("Normalize entity names: 4-M -> M4 threaded hole")
doc.add_paragraph("Results: 21 fixes across 6 files")

doc.add_heading("Step 3: QLoRA Training (03_qlora_train.py)", level=2)
doc.add_paragraph("Base model: Qwen2.5-7B-Instruct (preserved at /media/z/data/models/)")
doc.add_paragraph("Quantization: 4-bit nf4 (bitsandbytes)")
doc.add_paragraph("LoRA config: rank=16, alpha=32, target_modules=all-linear")
doc.add_paragraph("Training params: batch=1x2, max_seq=1024, lr=2e-4, epochs=3")
doc.add_paragraph("Trainable params: 40M / 4.4B (0.92%)")
doc.add_paragraph("Loss: 2.4 -> 0.1 (3 epochs)")

doc.add_heading("Step 4: Merge & Deploy (04_merge_and_deploy.sh)", level=2)
doc.add_paragraph("Merged LoRA adapter into full model at new path")
doc.add_paragraph("Base model preserved unchanged")
doc.add_paragraph("Updated P100 API config: MODEL_PATH, MAX_MEMORY, device_map")

# ── 3. Key Bug Fixes ──
doc.add_heading("3. Key Issues Resolved", level=1)

issues = [
    ("Test prompt mismatch",
     "Training used ~800 char instruction with entity/relation type lists and JSON format rules. "
     "Test inference used ~50 char generic message. Model output natural language instead of JSON. "
     "Fixed by including full instruction in test_inference()."),
    ("P100 GPU config",
     "MAX_MEMORY={0:'12GiB',1:'12GiB'} + device_map='auto' caused 'CUDA error: invalid device ordinal'. "
     "Fixed to MAX_MEMORY={0:'15GiB'} + device_map='cuda:0' for single GPU."),
    ("Augmentation infinite loop",
     "Appending to samples list while iterating caused cascading augmentation. "
     "Fixed by iterating over snapshot, collecting new samples separately."),
    ("DXF rule engine vs LLM extraction",
     "Uploading DWGs via /ingest/file goes through DXF rule engine which extracts CAD primitives "
     "(SW_NOTE_0, SW_CENTERMARKSYMBOL_0) not knowledge. To use the fine-tuned model, "
     "use /ingest/text with .txt descriptions."),
]

for title_text, desc in issues:
    doc.add_heading(title_text, level=2)
    doc.add_paragraph(desc)

# ── 4. Integration ──
doc.add_heading("4. Knowledge Graph Integration", level=1)
doc.add_paragraph(
    "Call path: User .txt description -> KG /api/v1/ingest/text -> "
    "LLMExtractor -> P100 fine-tuned model (http://10.117.29.24:5200/v1) -> Neo4j"
)
doc.add_paragraph(
    "The /chat UI at http://localhost:8000/chat queries Neo4j for entities "
    "and relations extracted by the fine-tuned model."
)
doc.add_paragraph(
    "IMPORTANT: DWG files go through DXF rule engine (CAD primitives). "
    "To use the fine-tuned model, upload .txt descriptions via batch_upload_text.py."
)

# ── 5. File Locations ──
doc.add_heading("5. Key Files", level=1)

files = [
    ("Windows local", "E:\\Knowledge Graph_robot\\scripts\\finetune\\"),
    ("P100 server", "/data/finetune/"),
    ("Base model", "/media/z/data/models/Qwen2.5-7B-Instruct (unchanged)"),
    ("Merged model", "/data/finetune/output/qwen2.5-7b-kg-robot-merged/"),
    ("Training data", "data/handcrafted_examples.json"),
    ("Labeling script", "enhanced_label.py"),
    ("Cleaning script", "clean_labeled.py"),
    ("Training script", "03_qlora_train.py"),
    ("Merge script", "04_merge_and_deploy.sh"),
    ("Text upload script", "batch_upload_text.py"),
    ("Desktop batch files", "04_enhanced_label.bat through 08_upload_text.bat"),
]

for name, path in files:
    p = doc.add_paragraph()
    p.add_run(f"{name}: ").bold = True
    p.add_run(path)

# ── 6. Next Steps ──
doc.add_heading("6. To Complete", level=1)
doc.add_paragraph("1. Confirm P100 API is running: curl http://10.117.29.24:5200/health")
doc.add_paragraph("2. Run batch_upload_text.py to re-ingest .txt descriptions through LLM extraction")
doc.add_paragraph("3. Verify KG Q&A at http://localhost:8000/chat returns correct answers")
doc.add_paragraph("4. Optionally clear old DXF rule engine data from Neo4j")

# ── Save ──
output = Path.home() / "Desktop" / "QLoRA_Fine-tuning_Report_2026-05-22.docx"
doc.save(str(output))
print(f"Saved: {output}")
