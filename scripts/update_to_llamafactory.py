#!/usr/bin/env python3
"""Update v2 docx: replace old QLoRA training with LLaMA-Factory workflow"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

INPUT = r"C:\Users\Knightz\Desktop\Qwen2.5-7B_QLoRA微调完整指南_v2.docx"
OUTPUT = r"C:\Users\Knightz\Desktop\Qwen2.5-7B_QLoRA微调完整指南_v3.docx"

doc = Document(INPUT)

# Find and update key paragraphs to reflect LLaMA-Factory
for i, p in enumerate(doc.paragraphs):
    # Update old training references
    if "03_qlora_train.py" in p.text:
        # Replace inline code references
        for run in p.runs:
            if "03_qlora_train.py" in run.text:
                run.text = run.text.replace("03_qlora_train.py", "LLaMA-Factory (train_config.yaml)")
    if "04_merge_and_deploy.sh" in p.text:
        for run in p.runs:
            if "04_merge_and_deploy.sh" in run.text:
                run.text = run.text.replace("04_merge_and_deploy.sh", "run_llamafactory.sh / merge_lora.py")

# ── Update heading "步骤3：QLoRA 训练 (03_qlora_train.py)" ──
for p in doc.paragraphs:
    if "QLoRA" in p.text and "03_qlora_train" in p.text:
        for run in p.runs:
            run.text = run.text.replace("03_qlora_train.py", "LLaMA-Factory CLI")
        break

# ── Add LLaMA-Factory section after the existing Step 3 content ──
# Find paragraph with "QLoRA 配置" or similar to insert after
insert_after = None
for i, p in enumerate(doc.paragraphs):
    if "Loss: 2.4" in p.text or "0.92%" in p.text:
        insert_after = p
        break

if insert_after:
    # Insert LLaMA-Factory training details
    # We need to find the parent element and insert after the found paragraph
    parent = insert_after._element.getparent()
    idx = list(parent).index(insert_after._element)

    def add_para(text, bold_prefix=None, style=None):
        from docx.oxml.ns import qn
        new_p = doc.add_paragraph()
        if bold_prefix:
            r = new_p.add_run(bold_prefix)
            r.bold = True
            new_p.add_run(text)
        else:
            new_p.add_run(text)
        return new_p._element

    def add_heading(text, level):
        from lxml import etree
        from docx.oxml.ns import qn
        h = doc.add_heading(text, level=level)
        return h._element

    elements = []

    # LLaMA-Factory overview
    h1 = doc.add_heading("LLaMA-Factory 训练详情（2026-05-22 更新）", level=2)
    elements.append(h1._element)

    elements.append(add_para(
        "训练框架已从自定义 transformers Trainer 迁移到 LLaMA-Factory v0.9.3。"
        "LLaMA-Factory 封装了模型加载、LoRA 配置、训练循环和导出合并等步骤，"
        "使用 YAML 配置文件和 CLI 命令行即可完成全部操作。"
    ))

    h2 = doc.add_heading("安装", level=3)
    elements.append(h2._element)
    elements.append(add_para(
        "conda activate qwen-api\n"
        "pip install llamafactory -i https://pypi.tuna.tsinghua.edu.cn/simple"
    ))

    h3 = doc.add_heading("YAML 配置文件 (train_config.yaml)", level=3)
    elements.append(h3._element)
    elements.append(add_para("关键参数："))
    config_items = [
        "model_name_or_path: /media/z/data/models/Qwen2.5-7B-Instruct",
        "dataset: kg_robot（在 dataset_info.json 中注册）",
        "template: qwen",
        "finetuning_type: lora",
        "stage: sft",
        "lora_rank: 16, lora_alpha: 32, lora_dropout: 0.05, lora_target: all",
        "quantization_method: bitsandbytes, quantization_bit: 4",
        "per_device_train_batch_size: 1, gradient_accumulation_steps: 2",
        "learning_rate: 2.0e-4, lr_scheduler_type: cosine",
        "num_train_epochs: 3, cutoff_len: 1024",
        "fp16: true, bf16: false（P100 不支持 bf16）",
        "optim: paged_adamw_8bit",
    ]
    for item in config_items:
        elements.append(add_para(f"  - {item}"))

    h4 = doc.add_heading("训练命令", level=3)
    elements.append(h4._element)
    elements.append(add_para(
        "CUDA_VISIBLE_DEVICES=0 llamafactory-cli train /data/finetune/train_config.yaml"
    ))
    elements.append(add_para(
        "训练结果: 24 条数据, 3 epochs, ~6 分钟, Loss 2.4 → 0.1, checkpoint-6"
    ))

    h5 = doc.add_heading("合并导出", level=3)
    elements.append(h5._element)
    elements.append(add_para(
        "LLaMA-Factory 的 llamafactory-cli export 在 P100 上 OOM，改用 PEFT 直接合并："
    ))
    elements.append(add_para(
        "python merge_lora.py  # 加载 fp16 模型到 GPU → merge_and_unload() → 直接 GPU 保存"
    ))
    elements.append(add_para(
        "合并模型: /data/finetune/output/qwen2.5-7b-kg-robot-merged-v2/"
    ))

    h6 = doc.add_heading("一键脚本", level=3)
    elements.append(h6._element)
    elements.append(add_para(
        "bash /data/finetune/run_llamafactory.sh  # 检查 → 训练 → 导出 → 部署，全自动"
    ))

    # Insert all elements after the found paragraph
    for elem in reversed(elements):
        parent.insert(idx + 1, elem)

# ── Update file location references ──
for p in doc.paragraphs:
    if "03_qlora_train.py" in p.text or "04_merge_and_deploy.sh" in p.text:
        for run in p.runs:
            if "03_qlora_train.py" in run.text:
                run.text = run.text.replace("03_qlora_train.py", "train_config.yaml (LLaMA-Factory)")
            if "04_merge_and_deploy.sh" in run.text:
                run.text = run.text.replace("04_merge_and_deploy.sh", "run_llamafactory.sh")

# ── Update training results section ──
for p in doc.paragraphs:
    if "03_qlora_train.py" in p.text:
        for run in p.runs:
            run.text = run.text.replace("（合并部署脚本）", "（一键训练+导出+部署）")

# Save
doc.save(OUTPUT)
print(f"Saved: {OUTPUT}")
