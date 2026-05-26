#!/usr/bin/env python3
"""生成 QLoRA 微调操作手册 .docx"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# ── 样式 ──
style = doc.styles['Normal']
font = style.font
font.name = 'Microsoft YaHei'
font.size = Pt(10.5)

# ── 标题 ──
title = doc.add_heading('Qwen2.5-7B QLoRA 微调操作手册', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run('工业机器人知识图谱抽取任务 | P100 16GB 单卡').bold = True
doc.add_paragraph('日期: 2026-05-21').alignment = WD_ALIGN_PARAGRAPH.CENTER

# ═══════════════════════ 1 ═══════════════════════
doc.add_heading('1. 架构概览', level=1)
doc.add_paragraph('')
doc.add_heading('Windows 本地 (开发 / 数据准备)', level=2)
doc.add_paragraph('路径: E:\\Knowledge Graph_robot\\scripts\\finetune\\')
doc.add_paragraph('作用: 存放脚本、数据、配置文件')
doc.add_paragraph('运行脚本: 02_prepare_data.py, 06_auto_label.py, 07_batch_dwg_label.py')
doc.add_paragraph('依赖: ezdxf, requests')

doc.add_heading('P100 服务器 (训练 / 推理)', level=2)
doc.add_paragraph('IP: 10.117.29.24 (SSH: z@10.117.29.24)')
doc.add_paragraph('路径: /data/finetune/')
doc.add_paragraph('作用: 训练脚本执行、API 推理服务')
doc.add_paragraph('依赖: torch, transformers, peft, bitsandbytes, accelerate, datasets')
doc.add_paragraph('HF 镜像: export HF_ENDPOINT="https://hf-mirror.com"')
doc.add_paragraph('API 端点: http://10.117.29.24:5200/chat')

doc.add_heading('文件传输 (Windows <-> P100)', level=2)
doc.add_paragraph('Windows -> P100:  scp "E:/.../file" z@10.117.29.24:/data/finetune/')
doc.add_paragraph('P100 -> Windows:  scp z@10.117.29.24:/data/finetune/file "E:/..."')

# ═══════════════════════ 2 ═══════════════════════
doc.add_heading('2. 文件结构', level=1)

file_tree = """E:\\Knowledge Graph_robot\\scripts\\finetune\\          (Windows 本地仓库)
  ├── 01_setup_env.sh              P100 环境安装脚本
  ├── 02_prepare_data.py           训练数据生成/导出/验证
  ├── 03_qlora_train.py            核心 QLoRA 训练脚本
  ├── 04_merge_and_deploy.sh       LoRA adapter 合并部署
  ├── 05_test_adapter.py           微调效果对比测试
  ├── 06_auto_label.py             API 批量标注纯文本
  ├── 07_batch_dwg_label.py        DWG 一步到位标注
  ├── data/
  │   ├── handcrafted_examples.json  训练数据 (主文件, 50-100条)
  │   ├── auto_labeled.json          API 自动标注产出
  │   ├── raw_texts.txt              待标注纯文本 (每行一条)
  │   └── dwg/                       DWG源文件 + 同名 .txt 描述
  └── output/
      ├── final_adapter/             LoRA adapter (训练产出, ~160MB)
      └── checkpoint-*/              训练中间检查点

/data/finetune/                          (P100 服务器)
  ├── 03_qlora_train.py
  ├── 05_test_adapter.py
  ├── 06_auto_label.py
  ├── 04_merge_and_deploy.sh
  ├── data/handcrafted_examples.json
  └── output/
      ├── final_adapter/
      └── qwen2.5-7b-kg-robot-merged/   (合并后的完整模型)"""
doc.add_paragraph(file_tree)

# ═══════════════════════ 3 ═══════════════════════
doc.add_heading('3. 环境配置', level=1)

doc.add_heading('3.1 P100 环境安装', level=2)
doc.add_paragraph('# SSH 到 P100')
doc.add_paragraph('ssh z@10.117.29.24')
doc.add_paragraph('# 安装依赖 (如未装)')
doc.add_paragraph('cd /data/finetune && bash 01_setup_env.sh')
doc.add_paragraph('')
doc.add_paragraph('训练依赖包: torch, transformers==4.46.0, peft==0.13.0, accelerate==1.0.0, bitsandbytes==0.44.0, datasets==2.21.0')
doc.add_paragraph('辅助依赖: sentencepiece, scipy, tensorboard, fire, packaging')

doc.add_heading('3.2 关键环境变量', level=2)
doc.add_paragraph('export HF_ENDPOINT="https://hf-mirror.com"    # P100 上必须设, 否则连不上 HuggingFace')
doc.add_paragraph('CUDA_VISIBLE_DEVICES=0                        # 限制使用 GPU 0')
doc.add_paragraph('')
doc.add_paragraph('# API 推理服务环境变量')
doc.add_paragraph('MODEL_PATH=/media/z/data/models/Qwen2.5-7B-Instruct     # 在 qwen_api.py 中修改')

doc.add_heading('3.3 Windows 本地配置', level=2)
doc.add_paragraph('# ODA File Converter 路径 (07_batch_dwg_label.py 中)')
doc.add_paragraph('ODA_CONVERTER = r"E:\\ODA\\ODAFileConverter.exe"')
doc.add_paragraph('# 或设置环境变量: set ODA_CONVERTER_PATH=E:\\ODA\\ODAFileConverter.exe')
doc.add_paragraph('')
doc.add_paragraph('# 06/07 脚本中的 API 地址')
doc.add_paragraph('API_URL = "http://10.117.29.24:5200/v1/chat/completions"')

# ═══════════════════════ 4 ═══════════════════════
doc.add_heading('4. 训练数据格式 (重要)', level=1)
doc.add_paragraph('每条数据 JSON 对象含三个字段:')

format_text = """{
  "instruction": "你是一个专业的工业机器人领域知识图谱信息抽取专家...",
  "input": "FANUC M-20iA 是6轴工业机器人，负载20kg，臂展1853mm...",
  "output": "{\\"entities\\":[{\\"name\\":\\"FANUC\\",\\"type\\":\\"Manufacturer\\",...}],\\"relations\\":[{\\"source\\":{\\"name\\":\\"FANUC\\",\\"type\\":\\"Manufacturer\\"},\\"target\\":{\\"name\\":\\"M-20iA\\",\\"type\\":\\"Robot\\"},\\"relation_type\\":\\"manufactures\\",...}]}"
}"""
doc.add_paragraph(format_text)

doc.add_paragraph('')
doc.add_paragraph('预定义实体类型 (13种):')
doc.add_paragraph('Robot, Manufacturer, Component, Reducer, ServoMotor, Controller, Sensor, ApplicationScenario, Process, EndEffector, Standard, Material, Software')

doc.add_paragraph('预定义关系类型 (10种):')
doc.add_paragraph('manufactures, uses_reducer, uses_servo, uses_controller, uses_component, applied_in, contains, complies_with, performs_process, process_material')

doc.add_paragraph(''
                   '注意事项:', style='List Bullet')
doc.add_paragraph('output 必须是合法 JSON 字符串（用 json.dumps 序列化）', style='List Bullet')
doc.add_paragraph('实体类型必须用预定义列表中的', style='List Bullet')
doc.add_paragraph('数值属性提取为数字, 不带单位', style='List Bullet')
doc.add_paragraph('不要编造文本中没有的信息', style='List Bullet')
doc.add_paragraph('建议 50-100 条, 质量 > 数量', style='List Bullet')

# ═══════════════════════ 5 ═══════════════════════
doc.add_heading('5. 数据准备流程', level=1)

doc.add_heading('方式一: 手工标注 (推荐, 质量最高)', level=2)

steps_a = [
    ('Step 1: 生成模板', 'cd /d E:\\Knowledge Graph_robot\\scripts\\finetune\npython 02_prepare_data.py'),
    ('Step 2: 编辑数据', '打开 data/handcrafted_examples.json\n按第4节格式添加你的标注数据\ninput = DWG图纸文字描述 (1-2句话)\noutput = 正确的实体+关系 JSON'),
    ('Step 3: 上传 P100', 'scp "E:/Knowledge Graph_robot/scripts/finetune/data/handcrafted_examples.json" z@10.117.29.24:/data/finetune/data/'),
]
for title_text, detail in steps_a:
    doc.add_heading(title_text, level=3)
    doc.add_paragraph(detail)

doc.add_heading('方式二: API 自动标注纯文本', level=2)

steps_b = [
    ('Step 1: 准备文本', '编辑 data/raw_texts.txt, 每行一条待标注文本\n如: CHX-3底盘旋转涡轮箱，材质HT250，包含蜗杆轴...'),
    ('Step 2: 确认API在跑', 'P100上: curl -s http://127.0.0.1:5200/health'),
    ('Step 3: 上传脚本和数据', 'scp "E:/.../06_auto_label.py" z@10.117.29.24:/data/finetune/\nscp "E:/.../data/raw_texts.txt" z@10.117.29.24:/data/finetune/data/'),
    ('Step 4: 运行标注 (P100)', 'conda activate qwen-api\ncd /data/finetune\npython 06_auto_label.py --test    # 先测试一条\npython 06_auto_label.py           # 正式批量'),
    ('Step 5: 拉取结果', 'scp z@10.117.29.24:/data/finetune/data/auto_labeled.json "E:/.../data/"'),
    ('Step 6: 人工修正', '打开 auto_labeled.json\n逐条检查: 实体名对不对? 类型对不对? 关系反没反?\n改完保存, 合并到 handcrafted_examples.json'),
]
for title_text, detail in steps_b:
    doc.add_heading(title_text, level=3)
    doc.add_paragraph(detail)

doc.add_heading('方式三: DWG 批量处理', level=2)
doc.add_paragraph('注意: SolidWorks DWG 内部通常无有效文字, 默认 LLM 描述质量差。强烈建议每个 DWG 配同名 .txt 描述文件。')
steps_c = [
    ('Step 1: 放置文件', '把 DWG 放到 data/dwg/\n为每个 DWG 创建同名 .txt 描述\n如: CHX-3底盘旋转涡轮箱.DWG + CHX-3底盘旋转涡轮箱.txt'),
    ('Step 2: 生成描述', 'python 07_batch_dwg_label.py --desc-only    # 如 .txt 已存在则跳过'),
    ('Step 3: 抽取标注', 'python 07_batch_dwg_label.py               # 产出 data/auto_labeled.json'),
]
for title_text, detail in steps_c:
    doc.add_heading(title_text, level=3)
    doc.add_paragraph(detail)

doc.add_heading('数据合并命令', level=2)
doc.add_paragraph('# 将 auto_labeled.json 追加到 handcrafted_examples.json\npython -c "import json; f=open(\'data/handcrafted_examples.json\',\'r\',encoding=\'utf-8\'); old=json.load(f); f.close(); f=open(\'data/auto_labeled.json\',\'r\',encoding=\'utf-8\'); new=json.load(f); f.close(); [item.pop(\'_source_file\',None) for item in new]; old.extend(new); f=open(\'data/handcrafted_examples.json\',\'w\',encoding=\'utf-8\'); json.dump(old,f,ensure_ascii=False,indent=2); f.close(); print(\'Total: \'+str(len(old)))"')

# ═══════════════════════ 6 ═══════════════════════
doc.add_heading('6. 训练', level=1)

doc.add_heading('6.1 训练前检查', level=2)
doc.add_paragraph('# P100 上执行')
doc.add_paragraph('ps aux | grep qwen_api     # 找到 API 进程 PID')
doc.add_paragraph('kill -9 <PID>               # 杀掉释放显存')
doc.add_paragraph('nvidia-smi                  # 确认显存 ~0 MiB')
doc.add_paragraph('ls /data/finetune/data/handcrafted_examples.json    # 确认数据存在')
doc.add_paragraph('wc -l data/handcrafted_examples.json                # 检查数据量')

doc.add_heading('6.2 训练命令', level=2)
doc.add_paragraph('# 上传最新脚本')
doc.add_paragraph('scp "E:/Knowledge Graph_robot/scripts/finetune/03_qlora_train.py" z@10.117.29.24:/data/finetune/')
doc.add_paragraph('')
doc.add_paragraph('# P100 上')
doc.add_paragraph('conda activate qwen-api       # 或 finetune')
doc.add_paragraph('export HF_ENDPOINT="https://hf-mirror.com"')
doc.add_paragraph('cd /data/finetune')
doc.add_paragraph('CUDA_VISIBLE_DEVICES=0 python 03_qlora_train.py train')
doc.add_paragraph('')
doc.add_paragraph('# 自定义参数示例')
doc.add_paragraph('python 03_qlora_train.py train --num_epochs 5 --learning_rate 1e-4 --lora_r 32')

doc.add_heading('6.3 P100 调优参数', level=2)

params = [
    ('max_seq_length', '1024', 'CAD 文本 1024 够用'),
    ('per_device_batch_size', '1', '单卡最小 batch'),
    ('gradient_accumulation', '2', '等效 batch = 1x2 = 2'),
    ('lora_r', '16', 'rank 越大越强但越慢'),
    ('lora_alpha', '32', 'alpha = 2x rank'),
    ('learning_rate', '2e-4', '学习率'),
    ('num_epochs', '3', '训练轮数 (数据少可加)'),
    ('bnb_4bit_quant', 'nf4', '4-bit 量化类型'),
    ('use_nested_quant', 'False', 'P100 显存紧，不开'),
    ('gradient_checkpointing', 'True', '用计算换显存'),
    ('attn_implementation', 'sdpa', 'P100 用 SDPA'),
]

table = doc.add_table(rows=len(params)+1, cols=3)
table.style = 'Light Grid Accent 1'
for i, text in enumerate(['参数', '值', '说明']):
    table.rows[0].cells[i].text = text
for j, (name, val, desc) in enumerate(params):
    table.rows[j+1].cells[0].text = name
    table.rows[j+1].cells[1].text = val
    table.rows[j+1].cells[2].text = desc

doc.add_paragraph('')
doc.add_paragraph('P100 训练参考时间: 50条/3epochs ≈ 15分钟 | 100条/3epochs ≈ 30分钟')

doc.add_heading('6.4 追加训练', level=2)
doc.add_paragraph('# 有新数据时在已有 adapter 基础上继续训练')
doc.add_paragraph('python 03_qlora_train.py continue --adapter_path ./output/final_adapter --new_data new_samples.json')

# ═══════════════ 6.5 LLaMA-Factory ═══════════════
doc.add_heading('6.5 LLaMA-Factory 训练 (推荐)', level=2)
doc.add_paragraph(
    'LLaMA-Factory 是国内最流行的 LLM 微调框架（GitHub 40k+ Stars），'
    '提供 Web UI + CLI 训练方式，替代手写 Python 训练脚本。'
    '内置 Qwen2.5 支持、自动处理 chat template、一键切换 QLoRA/DoRA/全量微调。'
)

doc.add_heading('安装', level=3)
doc.add_paragraph(
    '# P100 上执行 (优先 pip，无需 git)\n'
    'pip install llamafactory\n\n'
    '# 或通过脚本安装 (含 git 镜像 fallback)\n'
    'bash /data/finetune/08_setup_llamafactory.sh\n\n'
    '# 验证\n'
    'llamafactory-cli version'
)

doc.add_heading('数据注册', level=3)
doc.add_paragraph(
    '# 数据格式保持不变 (instruction/input/output)\n'
    '# 通过 dataset_info.json 注册数据集\n'
    '{\n'
    '  "kg_robot": {\n'
    '    "file_name": "handcrafted_examples.json",\n'
    '    "columns": { "prompt": "instruction", "query": "input", "response": "output" }\n'
    '  }\n'
    '}\n\n'
    '# 上传到 P100\n'
    'scp "E:/.../data/handcrafted_examples.json" z@10.117.29.24:/data/finetune/data/\n'
    'scp "E:/.../data/dataset_info.json" z@10.117.29.24:/data/finetune/data/\n'
    'scp "E:/.../09_train_config.yaml" z@10.117.29.24:/data/finetune/'
)

doc.add_heading('训练', level=3)
doc.add_paragraph(
    '# 方式1: 命令行 (YAML配置)\n'
    'export HF_ENDPOINT="https://hf-mirror.com"\n'
    'cd /data/LLaMA-Factory  # 或任意目录 (pip安装时)\n'
    'llamafactory-cli train /data/finetune/09_train_config.yaml\n\n'
    '# 方式2: Web UI (浏览器操作)\n'
    'llamafactory-cli webui\n'
    '# 打开 http://10.117.29.24:7860 (需端口转发)\n\n'
    '# 导出一键部署\n'
    'llamafactory-cli export /data/finetune/export_config.yaml'
)

doc.add_heading('对比', level=3)

table2 = doc.add_table(rows=8, cols=3)
table2.style = 'Light Grid Accent 1'
for i, text in enumerate(['功能', '手写脚本 (03_qlora_train.py)', 'LLaMA-Factory']):
    table2.rows[0].cells[i].text = text
rows_data = [
    ('改参数', '编辑 Python 代码', '改 YAML 或 Web UI 点选'),
    ('加数据', '手动合 JSON 文件', '放文件 + 注册 dataset_info'),
    ('换方法', '改代码 (LoRA->DoRA)', 'YAML 改一行: finetuning_type: dora'),
    ('chat template', '手写 <|im_start|> 标签', '选 template: qwen 即可'),
    ('断点续训', '代码指定 checkpoint', 'YAML 配置 resume_from_checkpoint'),
    ('监控', 'tensorboard 另开', '内置 loss 曲线面板'),
    ('部署', '手写 merge 脚本', 'llamafactory-cli export 一键导出'),
]
for j, (a, b, c) in enumerate(rows_data):
    table2.rows[j+1].cells[0].text = a
    table2.rows[j+1].cells[1].text = b
    table2.rows[j+1].cells[2].text = c

doc.add_paragraph('')
doc.add_paragraph('结论: 训练建议用 LLaMA-Factory (方便)，手写脚本保留作为备用和流程理解参考。')

# ═══════════════════════ 7 ═══════════════════════
doc.add_heading('7. 验证与部署', level=1)

doc.add_heading('7.1 对比测试', level=2)
doc.add_paragraph('# 上传测试脚本')
doc.add_paragraph('scp "E:/Knowledge Graph_robot/scripts/finetune/05_test_adapter.py" z@10.117.29.24:/data/finetune/')
doc.add_paragraph('')
doc.add_paragraph('# P100 上执行 (注意: 先 kill qwen_api)')
doc.add_paragraph('kill <PID>')
doc.add_paragraph('cd /data/finetune')
doc.add_paragraph('python 05_test_adapter.py')
doc.add_paragraph('')
doc.add_paragraph('# 对比检查:')
doc.add_paragraph('1. 输出 JSON 格式是否正确 (entities + relations 数组)')
doc.add_paragraph('2. 实体类型是否在预定义列表中')
doc.add_paragraph('3. 关系字段是否是 source/target/relation_type')
doc.add_paragraph('4. 微调后不应比原始模型差')

doc.add_heading('7.2 合并部署', level=2)
doc.add_paragraph('# 满意效果后，P100 上执行')
doc.add_paragraph('cd /data/finetune')
doc.add_paragraph('bash 04_merge_and_deploy.sh')
doc.add_paragraph('')
doc.add_paragraph('# 完成后:')
doc.add_paragraph('1. 合并模型位置: /data/finetune/output/qwen2.5-7b-kg-robot-merged/')
doc.add_paragraph('2. 修改 /data/qwen_api.py 中 MODEL_PATH = ".../qwen2.5-7b-kg-robot-merged"')
doc.add_paragraph('3. 重启 API: conda activate qwen-api && python /data/qwen_api.py &')
doc.add_paragraph('4. 验证: curl -sN http://127.0.0.1:5200/chat -H "Content-Type: application/json" -d "{..."}')

# ═══════════════════════ 8 ═══════════════════════
doc.add_heading('8. 常用命令速查表', level=1)

doc.add_heading('文件传输', level=2)
cmds = [
    ('上传脚本', 'scp "E:/.../03_qlora_train.py" z@10.117.29.24:/data/finetune/'),
    ('上传数据', 'scp "E:/.../data/handcrafted_examples.json" z@10.117.29.24:/data/finetune/data/'),
    ('下载标注', 'scp z@10.117.29.24:/data/finetune/data/auto_labeled.json "E:/.../data/"'),
    ('上传目录', 'scp -r "E:/.../data/dwg/" z@10.117.29.24:/data/finetune/data/'),
]
for label_text, cmd in cmds:
    doc.add_paragraph(f'{label_text}:', style='List Bullet')
    doc.add_paragraph(cmd)

doc.add_heading('P100 日常运维', level=2)
cmds_p100 = [
    ('查看 GPU', 'nvidia-smi'),
    ('查看进程', 'ps aux | grep qwen_api'),
    ('杀进程', 'kill -9 <PID>'),
    ('启动 API', 'conda activate qwen-api && python /data/qwen_api.py &'),
    ('测试 API', 'curl -sN http://127.0.0.1:5200/chat -H "Content-Type: application/json" -d \'{"model":"qwen2.5-7b","messages":[{"role":"user","content":"你好"}],"max_tokens":32}\''),
    ('HF 镜像', 'export HF_ENDPOINT="https://hf-mirror.com"'),
    ('训练', 'CUDA_VISIBLE_DEVICES=0 python 03_qlora_train.py train'),
    ('LLaMA-Factory训练', 'llamafactory-cli train /data/finetune/09_train_config.yaml'),
    ('LLaMA-Factory Web', 'llamafactory-cli webui'),
    ('断点续训', 'python 03_qlora_train.py train --resume_from ./output/checkpoint-100'),
]
for label_text, cmd in cmds_p100:
    doc.add_paragraph(f'{label_text}:', style='List Bullet')
    doc.add_paragraph(cmd)

doc.add_heading('Windows 本地', level=2)
cmds_win = [
    ('切目录', 'cd /d E:\\Knowledge Graph_robot\\scripts\\finetune'),
    ('生模板', 'python 02_prepare_data.py'),
    ('标注文本', 'python 06_auto_label.py [--test]'),
    ('DWG批处理', 'python 07_batch_dwg_label.py [--desc-only] [--test]'),
    ('合并数据', '见第5节合并命令'),
]
for label_text, cmd in cmds_win:
    doc.add_paragraph(f'{label_text}:', style='List Bullet')
    doc.add_paragraph(cmd)

# ═══════════════════════ 9 ═══════════════════════
doc.add_heading('9. 完整工作流', level=1)

flow = [
    ('1. 收集原始文本', 'DWG 图纸 -> 人工写描述 .txt / 整理机器人参数表 / 从 Neo4j 导出'),
    ('2. 标注训练数据', '手工写 JSON 或 06_auto_label.py 自动 + 人工修正 -> handcrafted_examples.json'),
    ('3. 上传 P100', 'scp 脚本 + 数据到 10.117.29.24:/data/finetune/'),
    ('4. 训练', 'kill API -> CUDA_VISIBLE_DEVICES=0 python 03_qlora_train.py train\n或 llmfactory-cli train /data/finetune/09_train_config.yaml (LLaMA-Factory)'),
    ('5. 验证', 'python 05_test_adapter.py 对比原始 vs 微调效果'),
    ('6. 部署', 'bash 04_merge_and_deploy.sh -> 修改 qwen_api.py MODEL_PATH -> 重启 API'),
    ('7. 迭代', '每次发现 LLM 抽错 -> 把正确结果加入 -> 重新微调 -> 越来越准'),
]
for title_text, desc in flow:
    doc.add_heading(title_text, level=2)
    doc.add_paragraph(desc)

# ═══════════════════════ 10 ═══════════════════════
doc.add_heading('10. 故障排查', level=1)

issues = [
    ('CUDA OOM',
     '1. ps aux | grep qwen_api -> kill 掉 API\n'
     '2. 减小 max_seq_length (1024->768)\n'
     '3. 确认 CUDA_VISIBLE_DEVICES=0 只用单卡\n'
     '4. 确认 nvidia-smi 显存空闲'),
    ('HuggingFace 下载超时',
     'export HF_ENDPOINT="https://hf-mirror.com"\n'
     '检查 P100 外网连通性'),
    ('Training loss 不降',
     '增大 LR (2e-4->5e-4)\n'
     '检查数据质量 (output 是否全是正确JSON)\n'
     '增加数据量到 50+'),
    ('微调后反而变差',
     '数据量太少 <20条 (最常见原因)\n'
     '数据中有错误标注 (如 D系列 %%c)\n'
     '减小 lora_r 到 8, 增大 lora_dropout 到 0.1'),
    ('API 422 错误',
     'qwen_api.py 用 /chat 端点\n'
     '检查 JSON 是否含非法控制字符'),
    ('06_auto_label JSON 解析失败',
     '正常, 模型输出格式不固定\n'
     '拉回本地人工修正'),
    ('ODA DWG 转换失败',
     '确认 E:\\ODA\\ODAFileConverter.exe 存在\n'
     'DWG 版本太高? 尝试 ODA GUI 手动转'),
    ('SCP Host key verification failed',
     '首次连接需确认: ssh z@10.117.29.24\n'
     '或删除 known_hosts 中旧记录'),
]
for title_text, fix in issues:
    doc.add_heading(title_text, level=2)
    doc.add_paragraph(fix)

# ═══════════════════════ 保存 ═══════════════════════
desktop = os.path.join(os.environ['USERPROFILE'], 'Desktop')
output_path = os.path.join(desktop, 'QLoRA微调操作手册.docx')
doc.save(output_path)
print(f'已保存: {output_path}')
