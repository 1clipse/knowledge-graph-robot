from __future__ import annotations

import json
from typing import Any, Dict, List, Optional

from fastapi import APIRouter, HTTPException
from fastapi.responses import StreamingResponse
from loguru import logger
from pydantic import BaseModel, Field

from api.deps import neo4j_client
from config.settings import get_config
from graph.query import GraphQuery
from graph.rag_retriever import GraphRagRetriever
from rag.citation import CitationVerifier

router = APIRouter()

_retriever: Optional[GraphRagRetriever] = None


def _get_retriever() -> GraphRagRetriever:
    global _retriever
    if _retriever is None and neo4j_client is not None:
        _retriever = GraphRagRetriever(neo4j_client)
    return _retriever


class AskRequest(BaseModel):
    question: str = Field(..., description="用户问题")
    top_k: int = Field(default=5, description="检索相关实体数量")
    max_hops: int = Field(default=3, description="多跳推理跳数")


class Citation(BaseModel):
    marker: str = ""           # e.g. "P1"
    path_index: int = 0        # index into reasoning_paths
    snippet: str = ""          # the answer text around the citation


class AskResponse(BaseModel):
    status: str
    question: str
    answer: str = ""
    relevant_entities: List[Dict[str, Any]] = Field(default_factory=list)
    reasoning_paths: List[Dict[str, Any]] = Field(default_factory=list)
    citations: List[Citation] = Field(default_factory=list)
    context_used: str = ""


QA_SYSTEM_PROMPT = """你是一个工业机器人领域的知识问答助手。基于提供的知识图谱上下文信息，回答用户的问题。

上下文包含从知识图谱中检索到的"推理路径"（标记为P1、P2...），每条路径展示了实体之间的多跳关系链。

规则：
1. 优先使用提供的上下文信息回答，利用多跳关系链进行推理
2. 如果上下文信息不足以回答问题，明确说明
3. 回答要准确、专业、有条理，用中文
4. 涉及具体参数时请给出数值
5. 引用来源：只在陈述具体事实时标注，每条信息末尾标一次即可。例如"FANUC M-20iA的负载为20kg[P1]。"
6. 重要：每条路径最多引用1-2次，不要重复堆叠多个引用标记，不要在句子末尾堆砌[P1][P2]...[P10]
7. 如果无法从上下文中找到相关信息，不要编造引用"""

_CITATION_RE = None  # compiled lazily


def _get_citation_re():
    global _CITATION_RE
    if _CITATION_RE is None:
        _CITATION_RE = __import__("re").compile(r"\[P(\d+)\]")
    return _CITATION_RE


def _parse_citations(answer: str, reasoning_paths: List[Dict[str, Any]]) -> List[Citation]:
    """Extract [P1], [P2] markers from answer and map to reasoning_paths."""
    pat = _get_citation_re()
    matches = pat.finditer(answer)
    seen: set = set()
    citations: List[Citation] = []
    for m in matches:
        idx = int(m.group(1)) - 1
        if idx not in seen and 0 <= idx < len(reasoning_paths):
            seen.add(idx)
            start = max(0, m.start() - 40)
            end = min(len(answer), m.end() + 40)
            snippet = answer[start:end].strip()
            citations.append(Citation(
                marker=f"P{idx + 1}",
                path_index=idx,
                snippet=snippet,
            ))
    return citations


def _build_context(question: str, top_k: int, max_hops: int) -> tuple:
    """Delegates retrieval to GraphRagRetriever."""
    retriever = _get_retriever()
    retrieval = retriever.retrieve(question, top_k=top_k, max_hops=max_hops)

    # Optionally attach community context
    retrieval = _try_attach_communities(retrieval)

    # Convert scored paths back to the legacy path dict format
    paths = [sp.to_dict() for sp in retrieval.scored_paths]

    return retrieval.context_used, retrieval.search_results, paths


def _try_attach_communities(retrieval) -> Any:
    """Attach community summaries if space permits."""
    try:
        retriever = _get_retriever()
        retriever.attach_community_context(retrieval)
    except Exception as e:
        logger.debug(f"Community context skipped: {e}")
    return retrieval


@router.post("/ask", response_model=AskResponse)
async def ask_question(request: AskRequest) -> AskResponse:
    if neo4j_client is None:
        raise HTTPException(status_code=503, detail="Database not connected")

    context_used, relevant_entities, reasoning_paths = _build_context(
        request.question, request.top_k, request.max_hops
    )

    if not context_used:
        return AskResponse(
            status="success",
            question=request.question,
            answer="知识图谱中未找到相关信息。请尝试用更具体的术语提问，或先录入相关数据。",
        )

    try:
        from openai import AsyncOpenAI
        from extractors.llm_utils import llm_chat

        config = get_config()
        client = AsyncOpenAI(base_url=config.llm.base_url, api_key=config.llm.api_key)

        content, _, _ = await llm_chat(
            client=client,
            model=config.llm.model,
            messages=[
                {"role": "system", "content": QA_SYSTEM_PROMPT},
                {
                    "role": "user",
                    "content": f"知识图谱推理路径：\n{context_used}\n\n用户问题：{request.question}",
                },
            ],
            temperature=0.3,
            max_tokens=1024,
        )
        answer = content or ""
    except Exception as e:
        logger.error(f"LLM QA failed: {e}")
        answer = f"无法生成回答（LLM服务不可用）。相关知识路径：\n{context_used}"

    simplified_paths = [_simplify_path(p) for p in reasoning_paths]
    citations = _parse_citations(answer, simplified_paths)

    return AskResponse(
        status="success",
        question=request.question,
        answer=answer,
        relevant_entities=relevant_entities,
        reasoning_paths=simplified_paths,
        citations=citations,
        context_used=context_used,
    )


def _simplify_path(path: Dict[str, Any]) -> Dict[str, Any]:
    """Strip verbose properties from path nodes/edges for API response."""
    simple_nodes = []
    for n in path.get("nodes", []):
        props = n.get("properties", {})
        simple_nodes.append({
            "labels": n.get("labels", []),
            "name": props.get("name", ""),
            "key_props": {k: v for k, v in props.items() if k not in ("name", "file") and v},
        })
    simple_edges = []
    for e in path.get("edges", []):
        simple_edges.append({
            "type": e.get("type", ""),
            "start": e.get("start", ""),
            "end": e.get("end", ""),
        })
    return {"nodes": simple_nodes, "edges": simple_edges}


@router.post("/ask/stream")
async def ask_question_stream(request: AskRequest):
    if neo4j_client is None:
        raise HTTPException(status_code=503, detail="Database not connected")

    config = get_config()
    context_used, relevant_entities, reasoning_paths = _build_context(
        request.question, request.top_k, request.max_hops
    )

    if not context_used:
        async def empty_generate():
            yield f"data: {json.dumps({'type': 'error', 'message': '未找到相关信息'}, ensure_ascii=False)}\n\n"
        return StreamingResponse(empty_generate(), media_type="text/event-stream")

    from openai import AsyncOpenAI
    client = AsyncOpenAI(base_url=config.llm.base_url, api_key=config.llm.api_key)

    async def generate():
        yield f"data: {json.dumps({'type': 'meta', 'context': context_used, 'entities': len(relevant_entities), 'paths': len(reasoning_paths)}, ensure_ascii=False)}\n\n"

        try:
            response = await client.chat.completions.create(
                model=config.llm.model,
                messages=[
                    {"role": "system", "content": QA_SYSTEM_PROMPT},
                    {"role": "user", "content": f"知识图谱推理路径：\n{context_used}\n\n用户问题：{request.question}"},
                ],
                temperature=0.3,
                max_tokens=1024,
                stream=True,
            )

            async for chunk in response:
                for choice in chunk.choices:
                    if choice.delta and choice.delta.content:
                        yield f"data: {json.dumps({'type': 'token', 'content': choice.delta.content}, ensure_ascii=False)}\n\n"

            yield f"data: {json.dumps({'type': 'done'}, ensure_ascii=False)}\n\n"

        except Exception as e:
            logger.error(f"LLM stream failed: {e}")
            yield f"data: {json.dumps({'type': 'error', 'message': str(e)}, ensure_ascii=False)}\n\n"

    return StreamingResponse(
        generate(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no",
        },
    )


# ── Compare two entities ────────────────────────────────────


class CompareRequest(BaseModel):
    entity_a: str = Field(..., description="第一个实体名称")
    entity_b: str = Field(..., description="第二个实体名称")
    type_a: str = Field(default="", description="实体A类型(可选)")
    type_b: str = Field(default="", description="实体B类型(可选)")


class CompareResponse(BaseModel):
    status: str
    entity_a: Dict[str, Any] = Field(default_factory=dict)
    entity_b: Dict[str, Any] = Field(default_factory=dict)
    common_relations: List[Dict[str, Any]] = Field(default_factory=list)
    comparison: str = ""


class CompareExportRequest(BaseModel):
    """Full comparison data sent by frontend for DOCX/PDF export."""
    entity_a: Dict[str, Any] = Field(default_factory=dict)
    entity_b: Dict[str, Any] = Field(default_factory=dict)
    common_relations: List[Dict[str, Any]] = Field(default_factory=list)
    comparison: str = ""


COMPARE_PROMPT = """你是一个工业机器人领域的专家。请基于提供的两个实体的属性信息，生成结构化的对比分析。

要求：
1. 逐项对比两个实体的关键属性
2. 指出各自的优势和差异
3. 如果某些属性一方有而另一方没有，明确标注
4. 用中文输出，格式为Markdown"""


@router.get("/communities")
async def get_communities():
    """Return graph communities detected by Louvain algorithm."""
    if neo4j_client is None:
        raise HTTPException(status_code=503, detail="Database not connected")

    from graph.communities import CommunityManager
    cm = CommunityManager(neo4j_client)
    communities = cm.detect()

    try:
        from openai import AsyncOpenAI
        config = get_config()
        client = AsyncOpenAI(base_url=config.llm.base_url, api_key=config.llm.api_key)
        summaries = await cm.summarize(client, config.llm.model)
        for c in communities:
            c["summary"] = summaries.get(c["id"], "")
    except Exception as e:
        logger.warning(f"Community summaries failed: {e}")
        for c in communities:
            c["summary"] = ""

    return {
        "status": "success",
        "count": len(communities),
        "communities": communities,
    }


@router.post("/compare", response_model=CompareResponse)
async def compare_entities(request: CompareRequest) -> CompareResponse:
    if neo4j_client is None:
        raise HTTPException(status_code=503, detail="Database not connected")

    graph_query = GraphQuery(neo4j_client)

    # Fetch entity A
    entity_a = _fetch_entity(graph_query, request.entity_a, request.type_a)
    if not entity_a:
        raise HTTPException(status_code=404, detail=f"Entity '{request.entity_a}' not found")

    # Fetch entity B
    entity_b = _fetch_entity(graph_query, request.entity_b, request.type_b)
    if not entity_b:
        raise HTTPException(status_code=404, detail=f"Entity '{request.entity_b}' not found")

    # Find common relations (same relation type to same target type)
    common = _find_common_relations(graph_query, entity_a, entity_b)

    # Build comparison context
    context = _build_comparison_context(entity_a, entity_b, common)

    # LLM comparison
    try:
        from openai import AsyncOpenAI
        from extractors.llm_utils import llm_chat

        config = get_config()
        client = AsyncOpenAI(base_url=config.llm.base_url, api_key=config.llm.api_key)

        comparison, _, _ = await llm_chat(
            client=client,
            model=config.llm.model,
            messages=[
                {"role": "system", "content": COMPARE_PROMPT},
                {"role": "user", "content": f"请对比以下两个实体：\n\n{context}"},
            ],
            temperature=0.3,
            max_tokens=1024,
        )
    except Exception as e:
        logger.error(f"Compare LLM failed: {e}")
        comparison = f"LLM对比服务不可用。原始数据：\n{context}"

    return CompareResponse(
        status="success",
        entity_a={"name": entity_a["name"], "labels": entity_a["labels"], "properties": entity_a["properties"]},
        entity_b={"name": entity_b["name"], "labels": entity_b["labels"], "properties": entity_b["properties"]},
        common_relations=common,
        comparison=comparison or "",
    )


# ── Export helpers ──────────────────────────────────────

import io
import re as _re_md
from datetime import datetime as _dt
from urllib.parse import quote as _url_quote

_FONT_PATH = r"C:\Windows\Fonts\msyh.ttc"


def _make_filename(entity_a: Dict[str, Any], entity_b: Dict[str, Any], ext: str) -> str:
    now = _dt.now().strftime("%Y%m%d_%H%M%S")
    na = (entity_a.get("name") or "entity_a").replace(" ", "_")
    nb = (entity_b.get("name") or "entity_b").replace(" ", "_")
    return f"Compare_{na}_vs_{nb}_{now}.{ext}"


def _content_disposition(filename: str) -> str:
    safe = _re_md.sub(r'[^A-Za-z0-9_.-]+', '_', filename).strip('_') or 'comparison_report'
    return f"attachment; filename=\"{safe}\"; filename*=UTF-8''{_url_quote(filename)}"


# ── PDF Export ──────────────────────────────────────────


@router.post("/compare/export/pdf")
async def compare_export_pdf(request: CompareExportRequest):
    """Generate and download a .pdf comparison report."""
    try:
        pdf_bytes = _generate_comparison_pdf(
            entity_a=request.entity_a,
            entity_b=request.entity_b,
            common_relations=request.common_relations,
            comparison=request.comparison,
        )
    except Exception as e:
        logger.error(f"PDF generation failed: {e}")
        raise HTTPException(status_code=500, detail=f"PDF生成失败: {e}")

    from fastapi.responses import Response
    filename = _make_filename(request.entity_a, request.entity_b, "pdf")
    return Response(
        content=pdf_bytes,
        media_type="application/pdf",
        headers={"Content-Disposition": _content_disposition(filename)},
    )


def _generate_comparison_pdf(
    entity_a: Dict[str, Any],
    entity_b: Dict[str, Any],
    common_relations: List[Dict[str, Any]],
    comparison: str,
) -> bytes:
    from fpdf import FPDF

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=18)
    pdf.add_page()

    pdf.add_font("YaHei", fname=_FONT_PATH)
    pdf.add_font("YaHei", style="B", fname=_FONT_PATH)

    name_a = str(entity_a.get("name", "实体A"))
    name_b = str(entity_b.get("name", "实体B"))
    props_a = entity_a.get("properties", {}) or {}
    props_b = entity_b.get("properties", {}) or {}

    def _w(text, bold=False, size=10.5, align="L", h=6):
        pdf.set_font("YaHei", "B" if bold else "", size)
        pdf.multi_cell(0, h, text, align=align)
        pdf.set_x(pdf.l_margin)
        pdf.ln(1)

    def _section_heading(text, level=1):
        pdf.ln(4)
        size = 15 if level == 0 else (13 if level == 1 else 11)
        _w(text, bold=True, size=size)
        pdf.ln(2)

    # Title
    _w("实体对比分析报告", bold=True, size=18, align="C")
    _w(f"{name_a}  vs  {name_b}", size=10, align="C")
    _w(_dt.now().strftime("%Y-%m-%d %H:%M"), size=9, align="C")
    pdf.ln(6)

    # Section 1: Properties table
    _section_heading("一、基本信息对比")
    all_keys = list(dict.fromkeys(list(props_a.keys()) + list(props_b.keys())))
    if not all_keys:
        all_keys = ["name"]

    col_w = [50, 65, 65]
    pdf.set_font("YaHei", "B", 9)
    pdf.set_fill_color(219, 234, 254)
    for i, hdr in enumerate(["属性", name_a, name_b]):
        pdf.cell(col_w[i], 7, hdr, border=1, fill=True, align="C")
    pdf.ln()

    pdf.set_font("YaHei", "", 9)
    for key in all_keys:
        va = str(props_a.get(key, "—")) if props_a.get(key) is not None else "—"
        vb = str(props_b.get(key, "—")) if props_b.get(key) is not None else "—"
        for i, val in enumerate([key, va, vb]):
            pdf.cell(col_w[i], 7, val[:40], border=1, align="C" if i == 0 else "L")
        pdf.ln()
    pdf.ln(4)

    # Section 2: Common relations
    if common_relations:
        _section_heading("二、共同关系")
        rel_col = [46, 67, 67]
        pdf.set_font("YaHei", "B", 9)
        pdf.set_fill_color(219, 234, 254)
        for i, hdr in enumerate(["关系类型", f"{name_a} 关联", f"{name_b} 关联"]):
            pdf.cell(rel_col[i], 7, hdr, border=1, fill=True, align="C")
        pdf.ln()
        pdf.set_font("YaHei", "", 9)
        for cr in common_relations:
            vals = [
                cr.get("relation_type", ""),
                ", ".join(cr.get("entity_a_targets", [])[:3]),
                ", ".join(cr.get("entity_b_targets", [])[:3]),
            ]
            for i, val in enumerate(vals):
                pdf.cell(rel_col[i], 7, val[:32], border=1, align="L")
            pdf.ln()
        pdf.ln(4)

    # Section 3: AI comparison
    sec_num = "三" if common_relations else "二"
    _section_heading(f"{sec_num}、AI 对比分析")
    if comparison:
        pdf.set_font("YaHei", "", 10)
        for line in comparison.split("\n"):
            stripped = line.strip()
            if not stripped:
                pdf.ln(3)
                continue
            if stripped.startswith("### "):
                pdf.set_font("YaHei", "B", 10.5)
                pdf.multi_cell(0, 6, stripped[4:])
                pdf.set_x(pdf.l_margin)
                pdf.set_font("YaHei", "", 10)
            elif stripped.startswith("## "):
                pdf.set_font("YaHei", "B", 12)
                pdf.multi_cell(0, 7, stripped[3:])
                pdf.set_x(pdf.l_margin)
                pdf.set_font("YaHei", "", 10)
            elif stripped.startswith("# "):
                pdf.set_font("YaHei", "B", 13)
                pdf.multi_cell(0, 7, stripped[2:])
                pdf.set_x(pdf.l_margin)
                pdf.set_font("YaHei", "", 10)
            elif stripped.startswith("- "):
                pdf.multi_cell(0, 6, "  • " + stripped[2:])
                pdf.set_x(pdf.l_margin)
            else:
                parts = _re_md.split(r"(\*\*.+?\*\*)", stripped)
                line_text = "".join(
                    part[2:-2] if part.startswith("**") and part.endswith("**") else part
                    for part in parts
                )
                pdf.multi_cell(0, 6, line_text)
                pdf.set_x(pdf.l_margin)
    else:
        _w("（无分析内容）", size=10)

    pdf.ln(8)
    _w("— 工业机器人知识图谱系统自动生成 —", size=8, align="C")

    return bytes(pdf.output())


# ── DOCX Export ─────────────────────────────────────────


@router.post("/compare/export/docx")
async def compare_export_docx(request: CompareExportRequest):
    """Generate and download a .docx comparison report."""
    try:
        doc = _generate_comparison_docx(
            entity_a=request.entity_a,
            entity_b=request.entity_b,
            common_relations=request.common_relations,
            comparison=request.comparison,
        )
    except Exception as e:
        logger.error(f"DOCX generation failed: {e}")
        raise HTTPException(status_code=500, detail=f"DOCX生成失败: {e}")

    from fastapi.responses import Response

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    filename = _make_filename(request.entity_a, request.entity_b, "docx")
    return Response(
        content=buf.getvalue(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": _content_disposition(filename)},
    )


def _generate_comparison_docx(
    entity_a: Dict[str, Any],
    entity_b: Dict[str, Any],
    common_relations: List[Dict[str, Any]],
    comparison: str,
):
    """Build a structured .docx comparison report using python-docx."""
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)

    style = doc.styles["Normal"]
    style.font.size = Pt(10.5)
    style.paragraph_format.space_after = Pt(4)

    name_a = str(entity_a.get("name", "实体A"))
    name_b = str(entity_b.get("name", "实体B"))
    props_a = entity_a.get("properties", {}) or {}
    props_b = entity_b.get("properties", {}) or {}

    # Title
    title = doc.add_heading("实体对比分析报告", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_keep_with_next(title)
    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_p.add_run(_dt.now().strftime("%Y-%m-%d %H:%M")).font.size = Pt(9)
    _set_keep_with_next(date_p)
    doc.add_paragraph()

    # Section 1: Properties table
    h1 = doc.add_heading("一、基本信息对比", level=1)
    _set_keep_with_next(h1)
    all_keys = list(dict.fromkeys(list(props_a.keys()) + list(props_b.keys())))
    if not all_keys:
        all_keys = ["name"]

    table = doc.add_table(rows=len(all_keys) + 1, cols=3)
    table.style = "Light Grid Accent 1"
    _set_table_rows_keep_together(table)
    hdr = table.rows[0].cells
    hdr[0].text = "属性"
    hdr[1].text = name_a
    hdr[2].text = name_b
    for cell in hdr:
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
    for i, key in enumerate(all_keys):
        row = table.rows[i + 1]
        row.cells[0].text = key
        row.cells[1].text = str(props_a.get(key, "—")) if props_a.get(key) is not None else "—"
        row.cells[2].text = str(props_b.get(key, "—")) if props_b.get(key) is not None else "—"
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
    doc.add_paragraph()

    # Section 2: Common relations
    if common_relations:
        h2 = doc.add_heading("二、共同关系", level=1)
        _set_keep_with_next(h2)
        rel_table = doc.add_table(rows=len(common_relations) + 1, cols=3)
        rel_table.style = "Light Grid Accent 1"
        _set_table_rows_keep_together(rel_table)
        rel_hdr = rel_table.rows[0].cells
        rel_hdr[0].text = "关系类型"
        rel_hdr[1].text = f"{name_a} 关联"
        rel_hdr[2].text = f"{name_b} 关联"
        for cell in rel_hdr:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.bold = True
                    run.font.size = Pt(9)
        for i, cr in enumerate(common_relations):
            row = rel_table.rows[i + 1]
            row.cells[0].text = cr.get("relation_type", "")
            row.cells[1].text = ", ".join(cr.get("entity_a_targets", [])[:5])
            row.cells[2].text = ", ".join(cr.get("entity_b_targets", [])[:5])
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(9)
        doc.add_paragraph()

    # Section 3: AI comparison narrative
    if comparison:
        section_num = "三" if common_relations else "二"
        h3 = doc.add_heading(f"{section_num}、AI 对比分析", level=1)
        _set_keep_with_next(h3)
        for para_text in comparison.split("\n"):
            para_text = para_text.strip()
            if not para_text:
                continue
            if para_text.startswith("### "):
                sh = doc.add_heading(para_text[4:], level=3)
                _set_keep_with_next(sh)
            elif para_text.startswith("## "):
                sh = doc.add_heading(para_text[3:], level=2)
                _set_keep_with_next(sh)
            elif para_text.startswith("# "):
                sh = doc.add_heading(para_text[2:], level=1)
                _set_keep_with_next(sh)
            elif para_text.startswith("- "):
                p = doc.add_paragraph(para_text[2:], style="List Bullet")
                p.paragraph_format.keep_together = True
            else:
                p = doc.add_paragraph()
                p.paragraph_format.keep_together = True
                _apply_markdown_bold(p, para_text)

    # Footer
    doc.add_paragraph()
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_p.add_run("— 工业机器人知识图谱系统自动生成 —")
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

    return doc


def _set_keep_with_next(paragraph):
    """Prevent a paragraph from being the last on its page."""
    from lxml import etree
    from docx.oxml.ns import qn as _qn
    pPr = paragraph._p.get_or_add_pPr()
    kn = pPr.find(_qn("w:keepNext"))
    if kn is None:
        kn = etree.SubElement(pPr, _qn("w:keepNext"))
    kn.set(_qn("w:val"), "true")


def _set_table_rows_keep_together(table):
    """Prevent table rows from splitting across pages."""
    from lxml import etree
    from docx.oxml.ns import qn as _qn
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                pPr = p._p.get_or_add_pPr()
                kl = pPr.find(_qn("w:keepLines"))
                if kl is None:
                    kl = etree.SubElement(pPr, _qn("w:keepLines"))
                kl.set(_qn("w:val"), "true")


def _apply_markdown_bold(paragraph, text: str):
    """Convert **text** spans to bold runs within a paragraph."""
    for run in paragraph.runs:
        run.text = ""
    parts = _re_md.split(r"(\*\*.+?\*\*)", text)
    first = True
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            if first and paragraph.runs and paragraph.runs[0].text == "":
                paragraph.runs[0].text = part[2:-2]
                paragraph.runs[0].bold = True
            else:
                run = paragraph.add_run(part[2:-2])
                run.bold = True
        else:
            if first and paragraph.runs and paragraph.runs[0].text == "":
                paragraph.runs[0].text = part
            else:
                paragraph.add_run(part)
        first = False


def _fetch_entity(graph_query: GraphQuery, name: str, label: str = "") -> Dict[str, Any]:
    """Search for entity by name, optionally with type filter."""
    results = graph_query.fulltext_search(name, limit=5)
    for item in results:
        node = item.get("node", {})
        n_name = node.get("name", "")
        n_labels = item.get("labels", [])
        if n_name == name or name.lower() in n_name.lower():
            if not label or label in n_labels:
                return {
                    "name": n_name,
                    "labels": n_labels,
                    "properties": {k: v for k, v in node.items() if not k.startswith("_") and v is not None},
                }
    return {}


def _find_common_relations(graph_query: GraphQuery, entity_a: Dict, entity_b: Dict) -> List[Dict[str, Any]]:
    """Find relation types that both entities share (to the same type of target)."""
    try:
        label_a = entity_a["labels"][0] if entity_a["labels"] else ""
        label_b = entity_b["labels"][0] if entity_b["labels"] else ""
        na = entity_a["name"]
        nb = entity_b["name"]

        # Neighbors of A
        neighbors_a = graph_query.neighbors(label_a, na, limit=50)
        neighbors_b = graph_query.neighbors(label_b, nb, limit=50)

        # Build sets of (relation_type, target_label)
        a_rels = {}
        for n in neighbors_a:
            rel = n.get("relation_type", "")
            tgt = n.get("node", {}).get("name", "")
            if rel and tgt:
                a_rels.setdefault(rel, []).append(tgt)

        b_rels = {}
        for n in neighbors_b:
            rel = n.get("relation_type", "")
            tgt = n.get("node", {}).get("name", "")
            if rel and tgt:
                b_rels.setdefault(rel, []).append(tgt)

        common = []
        for rel_type in set(a_rels.keys()) & set(b_rels.keys()):
            common.append({
                "relation_type": rel_type,
                "entity_a_targets": a_rels[rel_type][:5],
                "entity_b_targets": b_rels[rel_type][:5],
            })
        return common
    except Exception:
        return []


def _build_comparison_context(entity_a: Dict, entity_b: Dict, common: List[Dict]) -> str:
    parts = []

    def fmt_entity(label: str, e: Dict) -> str:
        props = e.get("properties", {})
        prop_lines = "\n".join(f"  {k}: {v}" for k, v in props.items() if k != "name")
        return f"{label}: {e['name']} ({'/'.join(e.get('labels', []))})\n{prop_lines}"

    parts.append(fmt_entity("实体A", entity_a))
    parts.append("")
    parts.append(fmt_entity("实体B", entity_b))

    if common:
        parts.append("\n共同关系类型：")
        for c in common:
            parts.append(f"  {c['relation_type']}: A→{c['entity_a_targets']}, B→{c['entity_b_targets']}")

    return "\n".join(parts)
