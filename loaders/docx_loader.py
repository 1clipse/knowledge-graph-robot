"""Microsoft Word .docx file loader."""

from __future__ import annotations

import os
from typing import List

from loguru import logger


class DocxLoader:
    """Extract text content from .docx files via python-docx.

    Extracts:
    - Paragraph text
    - Table cell text (structured as TSV-like rows)
    - Header/footer text
    """

    def load(self, file_path: str) -> str:
        """Extract all readable text from a .docx file."""
        try:
            from docx import Document
        except ImportError:
            raise ImportError("python-docx is required for .docx support. Install with: pip install python-docx")

        doc = Document(file_path)
        parts: List[str] = []

        # Headers & footers
        for section in doc.sections:
            for header in (section.header, section.first_page_header, section.even_page_header):
                if header:
                    h_text = "\n".join(p.text for p in header.paragraphs if p.text.strip())
                    if h_text.strip():
                        parts.append(h_text)
            for footer in (section.footer, section.first_page_footer, section.even_page_footer):
                if footer:
                    f_text = "\n".join(p.text for p in footer.paragraphs if p.text.strip())
                    if f_text.strip():
                        parts.append(f_text)

        # Body paragraphs
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                parts.append(text)

        # Tables
        for table in doc.tables:
            rows: List[str] = []
            for row in table.rows:
                cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                rows.append(" | ".join(cells))
            if rows:
                parts.append("\n" + "\n".join(rows))

        text = "\n".join(parts)
        logger.info(f"DOCX loaded: {len(text)} chars from {os.path.basename(file_path)}")
        return text

    def load_and_chunk(self, file_path: str, chunk_size: int = 2000) -> List[str]:
        """Load and split text into chunks no larger than chunk_size."""
        text = self.load(file_path)
        if len(text) <= chunk_size:
            return [text]
        chunks: List[str] = []
        for i in range(0, len(text), chunk_size):
            chunks.append(text[i : i + chunk_size])
        return chunks
